"""Phase 2: Document Text Extraction with Caching"""
import logging
from typing import Optional
from sqlalchemy.orm import Session
from innovo_backend.shared.processing_cache import get_cached_document_text, store_document_text

logger = logging.getLogger(__name__)


def extract_document_text(
    file_bytes: bytes,
    file_content_hash: str,
    file_type: str,
    db: Optional[Session] = None,
) -> Optional[str]:
    if not file_bytes or not file_content_hash:
        return None

    if db:
        try:
            cached_text = get_cached_document_text(db, file_content_hash)
            if cached_text:
                logger.info(f"[CACHE REUSE] Using cached document text for content_hash={file_content_hash[:16]}...")
                return cached_text
        except Exception as cache_error:
            logger.warning(f"Cache lookup failed, proceeding with extraction: {str(cache_error)}")

    extracted_text = None

    try:
        if file_type.lower() == "pdf":
            extracted_text = _extract_pdf_text(file_bytes)
        elif file_type.lower() in ("docx", "doc"):
            extracted_text = _extract_docx_text(file_bytes)
        else:
            logger.warning(f"Unsupported document type: {file_type}")
            return None

        if extracted_text:
            logger.info(f"[PROCESSING] Document extraction completed: type={file_type}, length={len(extracted_text)} chars")
            if db:
                try:
                    store_document_text(db, file_content_hash, extracted_text)
                except Exception as cache_error:
                    logger.warning(f"Failed to store document text in cache: {str(cache_error)}")
            return extracted_text
        else:
            logger.warning(f"Document extraction returned no text for type={file_type}")
            return None

    except Exception as e:
        logger.error(f"Document extraction error: type={file_type}, error={str(e)}")
        return None


def _extract_pdf_text(file_bytes: bytes) -> Optional[str]:
    try:
        import PyPDF2
        from io import BytesIO
        pdf_reader = PyPDF2.PdfReader(BytesIO(file_bytes))
        text_content = []
        for page in pdf_reader.pages:
            text = page.extract_text()
            if text:
                text_content.append(text)
        combined_text = "\n\n".join(text_content)
        return combined_text.strip() if combined_text.strip() else None
    except ImportError:
        logger.error("PyPDF2 library not installed.")
        return None
    except Exception as e:
        logger.error(f"PDF extraction error: {str(e)}")
        return None


def _extract_docx_text(file_bytes: bytes) -> Optional[str]:
    try:
        from docx import Document as DocxDocument
        from io import BytesIO
        docx = DocxDocument(BytesIO(file_bytes))
        text_content = []
        for paragraph in docx.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text.strip())
        for table in docx.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_content.append(cell.text.strip())
        combined_text = "\n\n".join(text_content)
        return combined_text.strip() if combined_text.strip() else None
    except ImportError:
        logger.error("python-docx library not installed.")
        return None
    except Exception as e:
        logger.error(f"DOCX extraction error: {str(e)}")
        return None
