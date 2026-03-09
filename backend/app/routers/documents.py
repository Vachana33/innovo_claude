from fastapi import APIRouter, Depends, HTTPException, status, Response, Query
from sqlalchemy.orm import Session, defer, make_transient
from sqlalchemy.orm.attributes import flag_modified
from sqlalchemy.exc import ProgrammingError
from app.database import get_db
from app.models import Document, Company, User, FundingProgram
from app.schemas import DocumentResponse, DocumentUpdate, ChatRequest, ChatResponse, ChatConfirmationRequest, DocumentListItem
from app.dependencies import get_current_user
from app.template_resolver import get_template_for_document
from typing import List, Optional, Tuple, Dict, Any
from datetime import datetime, timezone
import os
import json
import re
import logging
import io
import traceback
from openai import OpenAI

logger = logging.getLogger(__name__)

router = APIRouter()

# ============================================================================
# ROLE SEPARATION ENFORCEMENT
# ============================================================================
# This module maintains strict separation between:
#
# 1. INITIAL GENERATION (used by /generate-content):
#    - Function: _generate_batch_content()
#    - Purpose: Creates content from scratch for empty sections
#    - Prompt: Assumes empty sections, focuses on creation
#
# 2. CHAT EDITING (used by /chat):
#    - Function: _generate_section_content()
#    - Purpose: Modifies existing section content
#    - Prompt: Assumes existing content, focuses on modification
#
# VERIFIED: No cross-calling exists:
# - /generate-content ONLY calls _generate_batch_content()
# - /chat ONLY calls _generate_section_content()
# ============================================================================

def _safe_get_document_by_id(document_id: int, db: Session) -> Optional[Document]:
    """
    Safely query Document by ID, handling missing chat_history column gracefully.
    Returns Document object or None if not found.
    """

    try:
        # First try normal query
        document = db.query(Document).filter(Document.id == document_id).first()
        if document:
            # Ensure chat_history is initialized
            if hasattr(document, 'chat_history') and document.chat_history is None:
                document.chat_history = []
            elif not hasattr(document, 'chat_history'):
                document.chat_history = []
        return document
    except ProgrammingError as e:

        # CRITICAL: Rollback transaction immediately
        db.rollback()

        # Check if error is about chat_history column not existing
        error_str = str(e.orig) if hasattr(e, 'orig') else str(e)
        if 'chat_history' in error_str.lower() or 'undefinedcolumn' in error_str.lower():
            logger.warning(f"chat_history column does not exist. Using workaround for document {document_id}")
            # Try using defer to exclude chat_history from query
            try:
                document = db.query(Document).options(defer(Document.chat_history)).filter(
                    Document.id == document_id
                ).first()
                if document:
                    # Set chat_history to empty list in memory
                    try:
                        document.chat_history = []
                    except Exception:  # noqa: B110
                        # Silently ignore if chat_history column doesn't exist (legacy compatibility)
                        pass
                return document
            except Exception:
                # If defer also fails, rollback and use raw SQL as last resort
                db.rollback()
                logger.warning("Defer approach also failed. Using raw SQL workaround.")
                from sqlalchemy import text
                try:
                    result = db.execute(
                        text("""
                            SELECT id, company_id, type, content_json, updated_at
                            FROM documents
                            WHERE id = :doc_id
                            LIMIT 1
                        """),
                        {"doc_id": document_id}
                    )
                    row = result.first()
                    if row:
                        # Parse content_json if it's a string
                        content_json = row[3]
                        if isinstance(content_json, str):
                            content_json = json.loads(content_json)

                        # Create Document object from raw SQL result
                        document = Document(
                            id=row[0],
                            company_id=row[1],
                            type=row[2],
                            content_json=content_json,
                            updated_at=row[4]
                        )
                        # Make it transient so SQLAlchemy doesn't try to track it
                        make_transient(document)
                        # Set chat_history in memory (not persisted, column doesn't exist)
                        document.chat_history = []
                        return document
                    return None
                except Exception as sql_error:
                    db.rollback()
                    logger.error(f"Raw SQL workaround also failed: {str(sql_error)}")
                    return None
        else:
            # Re-raise if it's a different ProgrammingError
            logger.error(f"Unexpected ProgrammingError: {str(e)}", exc_info=True)
            raise

@router.get(
    "/documents/by-id/{document_id}",
    response_model=DocumentResponse,
)
def get_document_by_id(
    document_id: int,
    db: Session = Depends(get_db),  # noqa: B008
    current_user: User = Depends(get_current_user),  # noqa: B008
):
    """
    Load an existing document by id. Used when opening a document from the list.
    Returns 404 if not found. Validates ownership via document's company.
    """
    document = _safe_get_document_by_id(document_id, db)
    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found",
        )
    company = db.query(Company).filter(
        Company.id == document.company_id,
        Company.user_email == current_user.email,
    ).first()
    if not company:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found",
        )
    if hasattr(document, "chat_history") and document.chat_history is None:
        document.chat_history = []
    elif not hasattr(document, "chat_history"):
        document.chat_history = []
    return document


@router.get(
    "/documents/{company_id}/vorhabensbeschreibung",
    response_model=DocumentResponse
)
def get_document(
    company_id: int,
    funding_program_id: Optional[int] = Query(None, description="Funding program ID"),
    template_id: Optional[str] = Query(None, description="User template ID (UUID)"),
    template_name: Optional[str] = Query(None, description="System template name (e.g., 'wtt_v1')"),
    title: Optional[str] = Query(None, description="Optional document title when creating"),
    db: Session = Depends(get_db),  # noqa: B008
    current_user: User = Depends(get_current_user)  # noqa: B008
):
    """
    Create a new document for a company (when funding_program_id is provided).
    Does not reuse existing documents; always creates a new row.
    Legacy: when funding_program_id is not provided, returns or creates the single legacy doc per company.
    """
    # Normalize optional query params (empty string -> None so template lookup is safe)
    template_id = (template_id.strip() if isinstance(template_id, str) else None) or None
    template_name = (template_name.strip() if isinstance(template_name, str) else None) or None
    doc_title = (title.strip() if isinstance(title, str) and title else None) or None

    # Verify company exists and belongs to current user
    company = db.query(Company).filter(
        Company.id == company_id,
        Company.user_email == current_user.email
    ).first()
    if not company:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Company not found"
        )

    # Get or create document
    # Handle case where chat_history column doesn't exist in database
    document = None
    chat_history_missing = False  # Track if we know chat_history column doesn't exist

    # Case 1: funding_program_id provided - always create a new document (no reuse)
    if funding_program_id:
        # Verify funding program exists and belongs to current user
        funding_program = db.query(FundingProgram).filter(
            FundingProgram.id == funding_program_id,
            FundingProgram.user_email == current_user.email
        ).first()
        if not funding_program:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Funding program not found"
            )

        # Create new document from template (never reuse existing)
        doc_template_id = None
        doc_template_name = None

        if template_id and str(template_id).strip():
            try:
                import uuid
                doc_template_id = uuid.UUID(template_id)
                logger.info(f"[TEMPLATE RESOLVER] Creating document with user template_id: {template_id}")
            except (ValueError, TypeError):
                raise HTTPException(
                    status_code=status.HTTP_400_BAD_REQUEST,
                    detail=f"Invalid template_id format: {template_id}"
                )
        elif template_name and str(template_name).strip():
            doc_template_name = template_name
            logger.info(f"[TEMPLATE RESOLVER] Creating document with system template_name: {template_name}")

        class TempDocument:
            def __init__(self, template_id, template_name):
                self.template_id = template_id
                self.template_name = template_name

        temp_document = TempDocument(doc_template_id, doc_template_name)

        try:
            template = get_template_for_document(temp_document, db, current_user.email)
        except ValueError as e:
            logger.error(f"[TEMPLATE RESOLVER] Failed to resolve template: {str(e)}")
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"Template not found or invalid: {str(e)}"
            ) from None
        except Exception as e:
            logger.error(f"[TEMPLATE RESOLVER] Unexpected error resolving template: {str(e)}", exc_info=True)
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"Template not found or invalid: {str(e)}"
            ) from e

        if not template or not isinstance(template, dict):
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Template not found"
            )

        sections = template.get("sections", []) if isinstance(template.get("sections"), list) else []

        for section in sections:
            if section.get("id") == "4.1" and section.get("type") == "milestone_table":
                pass
            elif "content" not in section:
                section["content"] = ""

        try:
            document = Document(
                company_id=company_id,
                funding_program_id=funding_program_id,
                type="vorhabensbeschreibung",
                template_id=doc_template_id,
                template_name=doc_template_name,
                title=doc_title,
                content_json={"sections": sections}
            )
            db.add(document)
            db.commit()
            db.refresh(document)
            template_info = f"template_id={document.template_id}" if document.template_id else f"template_name={document.template_name or 'default'}"
            logger.info(f"[TEMPLATE RESOLVER] Created document {document.id} from {template_info} for company {company_id}")
        except Exception as e:
            db.rollback()
            logger.error(f"Failed to create document from template: {str(e)}", exc_info=True)
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Failed to create document: {str(e)}"
            ) from e

    # Case 2: No funding_program_id provided - return legacy document
    else:
        try:
            # Return legacy document (funding_program_id=NULL) if exists
            document = db.query(Document).filter(
                Document.company_id == company_id,
                Document.funding_program_id.is_(None),
                Document.type == "vorhabensbeschreibung"
            ).first()
        except ProgrammingError as e:
            # CRITICAL: Rollback transaction immediately - it's in failed state after the error
            db.rollback()

            # Check if error is about chat_history or funding_program_id column not existing
            error_str = str(e.orig) if hasattr(e, 'orig') else str(e)
            if 'chat_history' in error_str.lower() or 'funding_program_id' in error_str.lower() or 'undefinedcolumn' in error_str.lower():
                chat_history_missing = True
                logger.warning(f"Column does not exist in database. Using workaround for company {company_id}")
                # Try using defer to exclude problematic columns from query
                try:
                    document = db.query(Document).options(defer(Document.chat_history)).filter(
                        Document.company_id == company_id,
                        Document.funding_program_id.is_(None),
                        Document.type == "vorhabensbeschreibung"
                    ).first()
                    if document:
                        # Set chat_history to empty list in memory (column doesn't exist in DB)
                        try:
                            document.chat_history = []
                        except Exception:  # noqa: B110
                            # Silently ignore if chat_history column doesn't exist (legacy compatibility)
                            pass
                except Exception as defer_error:
                    # If defer also fails, rollback and use raw SQL as last resort
                    db.rollback()
                    logger.warning(f"Defer approach also failed: {str(defer_error)}. Using raw SQL workaround.")
                    from sqlalchemy import text
                    try:
                        result = db.execute(
                            text("""
                                SELECT id, company_id, type, content_json, updated_at
                                FROM documents
                                WHERE company_id = :company_id AND type = :doc_type
                                LIMIT 1
                            """),
                            {"company_id": company_id, "doc_type": "vorhabensbeschreibung"}
                        )
                        row = result.first()
                        if row:
                            # Parse content_json if it's a string (PostgreSQL might return it as string or dict)
                            content_json = row[3]
                            if isinstance(content_json, str):
                                content_json = json.loads(content_json)

                            # Create Document object from raw SQL result
                            # Use make_transient to prevent SQLAlchemy from tracking it
                            document = Document(
                                id=row[0],
                                company_id=row[1],
                                type=row[2],
                                content_json=content_json,
                                updated_at=row[4]
                            )
                            # Make it transient so SQLAlchemy doesn't try to track it
                            # This prevents SQLAlchemy from trying to access chat_history when serializing
                            make_transient(document)
                            # Set chat_history in memory (not persisted, column doesn't exist)
                            document.chat_history = []
                    except Exception as sql_error:
                        db.rollback()
                        logger.error(f"Raw SQL workaround also failed: {str(sql_error)}")
                        document = None
            else:
                # Re-raise if it's a different ProgrammingError
                logger.error(f"Unexpected ProgrammingError: {str(e)}", exc_info=True)
                raise

        # Create empty legacy document if it doesn't exist
        if not document:
            # If we know chat_history is missing, skip ORM and use raw SQL directly
            if chat_history_missing:
                logger.warning(f"chat_history column does not exist. Creating document using raw SQL for company {company_id}")
                from sqlalchemy import text
                try:
                    # Check if funding_program_id column exists
                    try:
                        result = db.execute(
                            text("""
                                INSERT INTO documents (company_id, funding_program_id, type, content_json, updated_at)
                                VALUES (:company_id, :funding_program_id, :doc_type, :content_json, NOW())
                                RETURNING id, company_id, type, content_json, updated_at
                            """),
                            {
                                "company_id": company_id,
                                "funding_program_id": None,  # Legacy document
                                "doc_type": "vorhabensbeschreibung",
                                "content_json": json.dumps({"sections": []})
                            }
                        )
                    except Exception:
                        # funding_program_id column doesn't exist, use old schema
                        result = db.execute(
                            text("""
                                INSERT INTO documents (company_id, type, content_json, updated_at)
                                VALUES (:company_id, :doc_type, :content_json, NOW())
                                RETURNING id, company_id, type, content_json, updated_at
                            """),
                            {
                                "company_id": company_id,
                                "doc_type": "vorhabensbeschreibung",
                                "content_json": json.dumps({"sections": []})
                            }
                        )
                    row = result.first()
                    if row:
                        # Parse content_json if it's a string (PostgreSQL might return it as string or dict)
                        content_json = row[3]
                        if isinstance(content_json, str):
                            content_json = json.loads(content_json)

                        # Create Document object from inserted row
                        document = Document(
                            id=row[0],
                            company_id=row[1],
                            type=row[2],
                            content_json=content_json,
                            updated_at=row[4]
                        )
                        # Make it transient to avoid SQLAlchemy tracking issues
                        make_transient(document)
                        # Set chat_history in memory only
                        document.chat_history = []
                        # Commit the raw SQL insert
                        db.commit()
                    else:
                        raise HTTPException(
                            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                            detail="Failed to create document: no row returned"
                        )
                except Exception as sql_error:
                    db.rollback()
                    logger.error(f"Failed to create document using raw SQL: {str(sql_error)}", exc_info=True)
                    raise HTTPException(
                        status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                        detail=f"Failed to create document: {str(sql_error)}"
                    ) from None
            else:
                # Normal ORM creation path (chat_history column exists)
                try:
                    document = Document(
                        company_id=company_id,
                        funding_program_id=None,  # Legacy document
                        type="vorhabensbeschreibung",
                        content_json={"sections": []}
                        # chat_history is not set here - will be initialized after creation if column exists
                    )
                    db.add(document)
                    db.commit()
                    db.refresh(document)
                    logger.info(f"Created legacy document {document.id} for company {company_id}")
                except ProgrammingError as create_error:
                    # Check if error is about chat_history column not existing
                    error_str = str(create_error.orig) if hasattr(create_error, 'orig') else str(create_error)
                    if 'chat_history' in error_str.lower() or 'undefinedcolumn' in error_str.lower():
                        logger.warning(f"chat_history column does not exist. Creating document using raw SQL for company {company_id}")
                        db.rollback()
                        # Use raw SQL to insert without chat_history column
                        from sqlalchemy import text
                        try:
                            # Check if funding_program_id column exists
                            try:
                                result = db.execute(
                                    text("""
                                        INSERT INTO documents (company_id, funding_program_id, type, content_json, updated_at)
                                        VALUES (:company_id, :funding_program_id, :doc_type, :content_json, NOW())
                                        RETURNING id, company_id, type, content_json, updated_at
                                    """),
                                    {
                                        "company_id": company_id,
                                        "funding_program_id": None,  # Legacy document
                                        "doc_type": "vorhabensbeschreibung",
                                        "content_json": json.dumps({"sections": []})
                                    }
                                )
                            except Exception:
                                # funding_program_id column doesn't exist, use old schema
                                result = db.execute(
                                    text("""
                                        INSERT INTO documents (company_id, type, content_json, updated_at)
                                        VALUES (:company_id, :doc_type, :content_json, NOW())
                                        RETURNING id, company_id, type, content_json, updated_at
                                    """),
                                    {
                                        "company_id": company_id,
                                        "doc_type": "vorhabensbeschreibung",
                                        "content_json": json.dumps({"sections": []})
                                    }
                                )
                            row = result.first()
                            if row:
                                # Parse content_json if it's a string (PostgreSQL might return it as string or dict)
                                content_json = row[3]
                                if isinstance(content_json, str):
                                    content_json = json.loads(content_json)

                                # Create Document object from inserted row
                                document = Document(
                                    id=row[0],
                                    company_id=row[1],
                                    type=row[2],
                                    content_json=content_json,
                                    updated_at=row[4]
                                )
                                # Make it transient to avoid SQLAlchemy tracking issues
                                make_transient(document)
                                # Set chat_history in memory only
                                document.chat_history = []
                                # Commit the raw SQL insert
                                db.commit()
                            else:
                                raise HTTPException(
                                    status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                                    detail="Failed to create document: no row returned"
                                )
                        except Exception as sql_error:
                            db.rollback()
                            logger.error(f"Failed to create document using raw SQL: {str(sql_error)}", exc_info=True)
                            raise HTTPException(
                                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                                detail=f"Failed to create document: {str(sql_error)}"
                            ) from None
                    else:
                        # Re-raise if it's a different ProgrammingError
                        db.rollback()
                        logger.error(f"Unexpected ProgrammingError during document creation: {str(create_error)}", exc_info=True)
                        raise
                except Exception as e:
                    db.rollback()
                    logger.error(f"Failed to create document for company {company_id}: {str(e)}", exc_info=True)
                    raise HTTPException(
                        status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                        detail=f"Failed to create document: {str(e)}"
                    ) from e

    # Ensure document was resolved or created
    if document is None:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found"
        )

    # Ensure chat_history is initialized if null
    # Handle case where column might not exist yet (migration not run)
    try:
        # Check if chat_history attribute exists (column exists in DB)
        if hasattr(document, 'chat_history'):
            if document.chat_history is None:
                try:
                    document.chat_history = []
                    db.commit()
                    db.refresh(document)
                    logger.debug(f"Initialized chat_history for document {document.id}")
                except Exception as e:
                    logger.warning(f"Failed to initialize chat_history (column may not exist yet): {str(e)}")
                    db.rollback()
                    # Set to empty list in memory even if DB update fails
                    document.chat_history = []
        else:
            # Column doesn't exist - set to empty list in memory only
            logger.warning(f"chat_history column does not exist in database for document {document.id}")
            document.chat_history = []
    except Exception as e:
        # Fallback: if anything goes wrong, set to empty list in memory
        logger.warning(f"Error initializing chat_history for document {document.id}: {str(e)}")
        try:
            document.chat_history = []
        except Exception:  # noqa: B110
            # Silently ignore if chat_history column doesn't exist (legacy compatibility)
            # If we can't even set it in memory, continue without it
            pass

    return document

@router.get("/documents", response_model=List[DocumentListItem])
def list_documents(
    db: Session = Depends(get_db),  # noqa: B008
    current_user: User = Depends(get_current_user)  # noqa: B008
):
    """
    List all documents for the current user.
    Returns documents with company and funding program information.
    """
    try:
        # Query documents via company relationship - only documents where company belongs to user
        documents = db.query(Document).join(Company).filter(
            Company.user_email == current_user.email
        ).order_by(Document.updated_at.desc()).all()

        # Build response with company and funding program info
        result = []
        for doc in documents:
            # Get company name
            company = db.query(Company).filter(Company.id == doc.company_id).first()
            company_name = company.name if company else f"Company {doc.company_id}"

            # Get funding program title if exists
            funding_program_title = None
            if doc.funding_program_id:
                funding_program = db.query(FundingProgram).filter(
                    FundingProgram.id == doc.funding_program_id
                ).first()
                funding_program_title = funding_program.title if funding_program else None

            result.append(DocumentListItem(
                id=doc.id,
                company_id=doc.company_id,
                company_name=company_name,
                funding_program_id=doc.funding_program_id,
                funding_program_title=funding_program_title,
                type=doc.type,
                title=getattr(doc, "title", None),
                updated_at=doc.updated_at
            ))

        logger.info(f"Retrieved {len(result)} documents for user {current_user.email}")
        return result
    except Exception as e:
        logger.error(f"Error fetching documents for user {current_user.email}: {str(e)}", exc_info=True)
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to fetch documents: {str(e)}"
        ) from e


@router.delete("/documents/{document_id}", status_code=status.HTTP_204_NO_CONTENT)
def delete_document(
    document_id: int,
    db: Session = Depends(get_db),  # noqa: B008
    current_user: User = Depends(get_current_user)  # noqa: B008
):
    """
    Delete a document. Only allowed if the document's company belongs to the current user.
    """
    document = _safe_get_document_by_id(document_id, db)
    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found"
        )
    company = db.query(Company).filter(
        Company.id == document.company_id,
        Company.user_email == current_user.email
    ).first()
    if not company:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found"
        )
    try:
        db.delete(document)
        db.commit()
        logger.info(f"Deleted document {document_id} for user {current_user.email}")
    except Exception as e:
        db.rollback()
        logger.error(f"Failed to delete document {document_id}: {str(e)}", exc_info=True)
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Failed to delete document"
        ) from e


@router.put(
    "/documents/{document_id}",
    response_model=DocumentResponse
)
def update_document(
    document_id: int,
    document_data: DocumentUpdate,
    db: Session = Depends(get_db),  # noqa: B008
    current_user: User = Depends(get_current_user)  # noqa: B008
):
    """
    Update a document's content.
    Phase 2.6: Validates that section titles cannot be changed after headings_confirmed=True
    """

    document = _safe_get_document_by_id(document_id, db)
    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found"
        )

    # Verify company belongs to current user
    company = db.query(Company).filter(
        Company.id == document.company_id,
        Company.user_email == current_user.email
    ).first()
    if not company:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found"
        )

    # Phase 2.6: Validate section structure changes if headings are confirmed
    if hasattr(document, 'headings_confirmed') and document.headings_confirmed:
        old_sections = document.content_json.get("sections", [])
        new_sections = document_data.content_json.get("sections", [])

        # Create maps for quick lookup
        old_by_id = {s.get("id"): s for s in old_sections if isinstance(s, dict) and "id" in s}
        new_by_id = {s.get("id"): s for s in new_sections if isinstance(s, dict) and "id" in s}

        # Check for any title changes
        for section_id, new_section in new_by_id.items():
            if section_id in old_by_id:
                old_title = old_by_id[section_id].get("title", "")
                new_title = new_section.get("title", "")
                if old_title != new_title:
                    logger.warning(f"Attempted to rename section '{section_id}' from '{old_title}' to '{new_title}' after headings confirmation")
                    raise HTTPException(
                        status_code=status.HTTP_400_BAD_REQUEST,
                        detail="Headings are locked after confirmation. Section titles cannot be changed."
                    )

        # Check for new sections (section insertion)
        old_ids = set(old_by_id.keys())
        new_ids = set(new_by_id.keys())
        added_ids = new_ids - old_ids
        if added_ids:
            logger.warning(f"Attempted to add new sections {added_ids} after headings confirmation")
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Headings are locked after confirmation. New sections cannot be added."
            )

    # Update content
    document.content_json = document_data.content_json

    try:
        db.commit()
        db.refresh(document)
        return document
    except Exception as e:
        db.rollback()
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to update document: {str(e)}"
        ) from e


@router.post("/documents/{document_id}/confirm-headings", response_model=DocumentResponse)
def confirm_headings(
    document_id: int,
    db: Session = Depends(get_db),  # noqa: B008
    current_user: User = Depends(get_current_user)  # noqa: B008
):
    """
    Phase 2.6: Mark document headings as confirmed.
    This locks section titles from further changes.
    """
    document = _safe_get_document_by_id(document_id, db)
    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found"
        )

    # Verify company belongs to current user
    company = db.query(Company).filter(
        Company.id == document.company_id,
        Company.user_email == current_user.email
    ).first()
    if not company:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found"
        )

    # Set headings_confirmed to True
    if hasattr(document, 'headings_confirmed'):
        document.headings_confirmed = True
    else:
        # Fallback for databases that haven't run migration yet
        logger.warning(f"headings_confirmed column not found for document {document_id}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Headings confirmation not available. Please run database migration."
        )

    try:
        db.commit()
        db.refresh(document)
        logger.info(f"Headings confirmed for document {document_id}")
        return document
    except Exception as e:
        db.rollback()
        logger.error(f"Failed to confirm headings for document {document_id}: {str(e)}", exc_info=True)
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to confirm headings: {str(e)}"
        ) from e


# PDF style reference cache (extracted once, reused many times)
_pdf_style_reference_cache: Optional[str] = None


def _extract_pdf_style_reference(pdf_path: str) -> str:
    """
    Extract text from a PDF file to use as style reference.
    Returns a cleaned text sample (first 2000-3000 chars) that represents style, tone, and structure.
    If extraction fails, returns empty string (silent fallback).

    CRITICAL: This function must not raise exceptions - it must gracefully handle all errors.
    """
    try:
        import PyPDF2

        if not os.path.exists(pdf_path):
            logger.warning(f"PDF file not found: {pdf_path}")
            return ""

        text_content = []
        with open(pdf_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            # Extract text from first few pages (usually contains style examples)
            max_pages = min(3, len(pdf_reader.pages))
            for page_num in range(max_pages):
                page = pdf_reader.pages[page_num]
                text = page.extract_text()
                if text:
                    text_content.append(text)

            # Combine and clean text
            combined_text = "\n\n".join(text_content)

            # Take first 2500 characters as style sample
            style_sample = combined_text[:2500].strip()

            # Clean up excessive whitespace
            style_sample = re.sub(r'\s+', ' ', style_sample)
            style_sample = re.sub(r'\n\s*\n', '\n\n', style_sample)

            return style_sample

    except ImportError:
        logger.warning("PyPDF2 library not installed. PDF style references will not be available.")
        return ""
    except Exception as e:
        logger.warning(f"Failed to extract text from PDF {pdf_path}: {str(e)}")
        return ""


def _build_style_reference_text() -> str:
    """
    Build style reference text from both DIlico.pdf and Lagotec.pdf.
    Caches the result to avoid re-extracting on every call.
    Returns formatted text that can be inserted into prompts.
    If extraction fails, returns empty string (silent fallback).

    CRITICAL: This function must not raise exceptions - it must gracefully handle all errors.
    """
    global _pdf_style_reference_cache

    # Return cached result if available
    if _pdf_style_reference_cache is not None:
        return _pdf_style_reference_cache

    try:
        # Get the base directory (backend/app)
        base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        pdf_dir = os.path.join(base_dir, "app", "ai", "prompts", "vorhabensbeschreibung")

        dlico_path = os.path.join(pdf_dir, "DIlico.pdf")
        lagotec_path = os.path.join(pdf_dir, "Lagotec.pdf")

        dlico_text = _extract_pdf_style_reference(dlico_path)
        lagotec_text = _extract_pdf_style_reference(lagotec_path)

        # Build style reference section
        style_parts = []

        if dlico_text or lagotec_text:
            style_parts.append("WICHTIG - STILREFERENZEN:")
            style_parts.append("Sie haben Zugriff auf zwei Beispiel-Vorhabensbeschreibungen als PDF-Dateien (DIlico.pdf und Lagotec.pdf).")
            style_parts.append("Diese PDFs dienen AUSSCHLIESSLICH als Stilreferenzen für:")
            style_parts.append("- Ton und Formulierungsstil")
            style_parts.append("- Absatzlänge und narrative Tiefe")
            style_parts.append("- Strukturdichte")
            style_parts.append("- Formalisierungsgrad")
            style_parts.append("- Fördermittel-typische Struktur")
            style_parts.append("")
            style_parts.append("KRITISCH:")
            style_parts.append("- Kopieren Sie KEINEN Inhalt aus diesen PDFs")
            style_parts.append("- Paraphrasieren Sie KEINEN Inhalt aus diesen PDFs")
            style_parts.append("- Verwenden Sie KEINE Fakten aus diesen Dokumenten")
            style_parts.append("- Erwähnen Sie diese PDFs NICHT im generierten Text")
            style_parts.append("- Alle faktischen Inhalte müssen AUSSCHLIESSLICH aus den bereitgestellten Firmeninformationen stammen")
            style_parts.append("")

            if dlico_text:
                style_parts.append("STILBEISPIEL aus DIlico.pdf (NUR als Stilreferenz):")
                style_parts.append(dlico_text)
                style_parts.append("")

            if lagotec_text:
                style_parts.append("STILBEISPIEL aus Lagotec.pdf (NUR als Stilreferenz):")
                style_parts.append(lagotec_text)
                style_parts.append("")

            style_parts.append("Passen Sie Absatzlänge, narrative Dichte und professionellen Ton an den Stil der Beispiel-PDFs an.")
            style_parts.append("")

        result = "\n".join(style_parts)

        # Cache the result (even if empty)
        _pdf_style_reference_cache = result

        return result

    except Exception as e:
        logger.warning(f"Failed to build style reference text: {str(e)}")
        # Cache empty string to avoid repeated failures
        _pdf_style_reference_cache = ""
        return ""


def _format_company_context_for_prompt(
    company_profile: Optional[Dict[str, Any]],
    company_name: str,
    website_clean_text: Optional[str] = None,
    transcript_clean: Optional[str] = None,
    company_id: Optional[int] = None
) -> str:
    """
    Format company context for LLM prompts.
    
    Uses structured company_profile as PRIMARY factual source,
    enriched with cleaned website and transcript text for context.

    Args:
        company_profile: Structured company profile dict (PRIMARY source)
        company_name: Company name
        website_clean_text: Cleaned website text (contextual enrichment)
        transcript_clean: Cleaned transcript text (contextual enrichment)
        company_id: Company ID for logging (optional)

    Returns:
        Formatted company context string for prompt
    """
    context_parts = []
    
    # PRIMARY SOURCE: Structured company profile
    if company_profile:
        if company_id:
            logger.info(f"Using structured company_profile as PRIMARY source for company_id={company_id}")
        
        context_parts.append("=== PRIMÄRE FAKTENQUELLE (Strukturiertes Firmenprofil) ===")
        context_parts.append(f"Firmenname: {company_name}")

        if company_profile.get("industry"):
            context_parts.append(f"Branche: {company_profile['industry']}")

        if company_profile.get("products_or_services"):
            products = company_profile["products_or_services"]
            if isinstance(products, list) and products:
                products_str = ", ".join(products)
                context_parts.append(f"Produkte/Dienstleistungen: {products_str}")
            elif isinstance(products, str):
                context_parts.append(f"Produkte/Dienstleistungen: {products}")

        if company_profile.get("business_model"):
            context_parts.append(f"Geschäftsmodell: {company_profile['business_model']}")

        if company_profile.get("market"):
            context_parts.append(f"Zielmarkt: {company_profile['market']}")

        if company_profile.get("innovation_focus"):
            context_parts.append(f"Innovationsschwerpunkt: {company_profile['innovation_focus']}")

        if company_profile.get("company_size"):
            context_parts.append(f"Unternehmensgröße: {company_profile['company_size']}")

        if company_profile.get("location"):
            context_parts.append(f"Standort: {company_profile['location']}")
    else:
        if company_id:
            logger.warning(f"company_profile missing for company_id={company_id}, using name only")
        context_parts.append("=== PRIMÄRE FAKTENQUELLE ===")
        context_parts.append(f"Firmenname: {company_name}")
    
    # CONTEXTUAL ENRICHMENT: Cleaned texts
    if website_clean_text or transcript_clean:
        context_parts.append("\n=== KONTEXTUELLE ERGÄNZUNG ===")
        
        if website_clean_text:
            # Smart truncation for cleaned website text
            MAX_TEXT_LENGTH = 30000
            if len(website_clean_text) > MAX_TEXT_LENGTH:
                first_part = website_clean_text[:int(MAX_TEXT_LENGTH * 0.6)]
                last_part = website_clean_text[-int(MAX_TEXT_LENGTH * 0.4):]
                website_clean_text = f"{first_part}\n\n[... Inhalt gekürzt ...]\n\n{last_part}"
            context_parts.append(f"Website-Inhalt (bereinigt):\n{website_clean_text}")
        
        if transcript_clean:
            # Smart truncation for cleaned transcript
            MAX_TEXT_LENGTH = 30000
            if len(transcript_clean) > MAX_TEXT_LENGTH:
                first_part = transcript_clean[:int(MAX_TEXT_LENGTH * 0.6)]
                last_part = transcript_clean[-int(MAX_TEXT_LENGTH * 0.4):]
                transcript_clean = f"{first_part}\n\n[... Inhalt gekürzt ...]\n\n{last_part}"
            context_parts.append(f"Besprechungsprotokoll (bereinigt):\n{transcript_clean}")
    
    return "\n".join(context_parts)


def _split_sections_into_batches(sections: List[dict], batch_size: int = 4) -> List[List[dict]]:
    """
    Split sections into batches of 3-5 headings for chunked generation.
    Default batch_size is 4, but will vary between 3-5 to balance efficiency and reliability.
    """
    batches = []
    current_batch = []

    for section in sections:
        current_batch.append(section)
        if len(current_batch) >= batch_size:
            batches.append(current_batch)
            current_batch = []
            # Vary batch size slightly (3-5) for better distribution
            batch_size = 3 if batch_size == 5 else 5 if batch_size == 3 else 4

    # Add remaining sections as final batch
    if current_batch:
        batches.append(current_batch)

    return batches


def _generate_batch_content(
    client: OpenAI,
    batch_sections: List[dict],
    company_name: str,
    company_profile: Optional[Dict[str, Any]] = None,
    website_clean_text: Optional[str] = None,
    transcript_clean: Optional[str] = None,
    company_id: Optional[int] = None,
    funding_program_rules: Optional[Dict[str, Any]] = None,
    style_profile: Optional[Dict[str, Any]] = None,
    max_retries: int = 2
) -> dict:
    """
    ROLE: INITIAL GENERATION

    Creates section content from scratch for empty or new sections.
    Used ONLY during first draft generation via /generate-content endpoint.

    This function:
    - Assumes sections are empty or need initial content
    - Focuses on creation and expansion
    - Can be creative and comprehensive
    - Generates content based on company data and style references

    This function must NOT:
    - Be used for chat-based editing
    - Modify existing section content
    - Be called from /chat endpoint

    Returns a dictionary mapping section_id to generated content.
    Implements retry logic with strict JSON validation.
    """
    # Build headings list for this batch (exclude milestone tables)
    headings_list = []
    section_ids = []
    for section in batch_sections:
        # Skip milestone tables - they should not be AI-generated
        if section.get('type') == 'milestone_table':
            continue
        section_id = section.get('id', '')
        section_title = section.get('title', '')
        # Remove numbering prefix from title
        clean_title = re.sub(r'^[\d.]+\.\s*', '', section_title)
        headings_list.append(f"{section_id}. {clean_title}")
        section_ids.append(section_id)

    headings_text = "\n".join(headings_list)

    # IMPORTANT: This prompt is for INITIAL CONTENT GENERATION only.
    # It assumes empty sections and focuses on creation.
    # Do NOT reuse this prompt for chat-based editing.
    # For editing existing content, use _generate_section_content() instead.

    # ============================================
    # PROMPT STRUCTURE: Rules → Company → Style → Task
    # ============================================

    # 1. RULES SECTION (from funding program guidelines)
    rules_section = ""
    if funding_program_rules:
        rules_parts = []
        if funding_program_rules.get("eligibility_rules"):
            rules_parts.append("Berechtigungskriterien:\n" + "\n".join(f"- {r}" for r in funding_program_rules["eligibility_rules"]))
        if funding_program_rules.get("required_sections"):
            rules_parts.append("Erforderliche Abschnitte:\n" + "\n".join(f"- {r}" for r in funding_program_rules["required_sections"]))
        if funding_program_rules.get("forbidden_content"):
            rules_parts.append("Verbotene Inhalte:\n" + "\n".join(f"- {r}" for r in funding_program_rules["forbidden_content"]))
        if funding_program_rules.get("formal_requirements"):
            rules_parts.append("Formale Anforderungen:\n" + "\n".join(f"- {r}" for r in funding_program_rules["formal_requirements"]))
        if funding_program_rules.get("evaluation_criteria"):
            rules_parts.append("Bewertungskriterien:\n" + "\n".join(f"- {r}" for r in funding_program_rules["evaluation_criteria"]))
        if funding_program_rules.get("funding_limits"):
            rules_parts.append("Fördergrenzen:\n" + "\n".join(f"- {r}" for r in funding_program_rules["funding_limits"]))
        if funding_program_rules.get("deadlines"):
            rules_parts.append("Fristen:\n" + "\n".join(f"- {r}" for r in funding_program_rules["deadlines"]))
        if funding_program_rules.get("important_notes"):
            rules_parts.append("Wichtige Hinweise:\n" + "\n".join(f"- {r}" for r in funding_program_rules["important_notes"]))
        
        if rules_parts:
            rules_section = "=== 1. FÖRDERRICHTLINIEN UND REGELN ===\n\n" + "\n\n".join(rules_parts) + "\n\n"

    # 2. COMPANY SOURCE SECTION (primary: company_profile, enrichment: cleaned texts)
    company_context = _format_company_context_for_prompt(
        company_profile=company_profile,
        company_name=company_name,
        website_clean_text=website_clean_text,
        transcript_clean=transcript_clean,
        company_id=company_id
    )
    company_section = f"=== 2. FIRMENINFORMATIONEN (FAKTENQUELLE) ===\n\n{company_context}\n\n"

    # 3. STYLE GUIDE SECTION (from AlteVorhabensbeschreibungStyleProfile)
    style_section = ""
    if style_profile:
        style_parts = []
        
        if style_profile.get("structure_patterns"):
            patterns = style_profile["structure_patterns"]
            if isinstance(patterns, list) and patterns:
                style_parts.append("Strukturmuster:\n" + "\n".join(f"- {p}" for p in patterns))
        
        if style_profile.get("tone_characteristics"):
            tone = style_profile["tone_characteristics"]
            if isinstance(tone, list) and tone:
                style_parts.append("Ton und Charakteristik:\n" + "\n".join(f"- {t}" for t in tone))
        
        if style_profile.get("writing_style_rules"):
            rules = style_profile["writing_style_rules"]
            if isinstance(rules, list) and rules:
                style_parts.append("Schreibstil-Regeln:\n" + "\n".join(f"- {r}" for r in rules))
        
        if style_profile.get("storytelling_flow"):
            flow = style_profile["storytelling_flow"]
            if isinstance(flow, list) and flow:
                style_parts.append("Erzählstruktur und Flow:\n" + "\n".join(f"- {f}" for f in flow))
        
        if style_profile.get("common_section_headings"):
            headings = style_profile["common_section_headings"]
            if isinstance(headings, list) and headings:
                style_parts.append("Typische Abschnittsüberschriften:\n" + "\n".join(f"- {h}" for h in headings))
        
        if style_parts:
            style_section = "=== 3. STIL-LEITFADEN ===\n\n" + "\n\n".join(style_parts) + "\n\n"
            style_section += "WICHTIG: Folgen Sie diesen Stilrichtlinien STRENG bei der Generierung.\n"
            style_section += "Passen Sie Ton, Struktur, Satzlänge und Erzählweise an diese Vorgaben an.\n\n"
    else:
        logger.warning("No style profile available, using default style guidelines")
        style_section = "=== 3. STIL-LEITFADEN ===\n\n"
        style_section += "- Verwenden Sie formelle Fördermittel-/Geschäftssprache\n"
        style_section += "- Professioneller, überzeugender Ton\n"
        style_section += "- Klare Absatzstruktur\n\n"

    # 4. GENERATION TASK
    task_section = f"""=== 4. GENERIERUNGSAUFGABE ===

Zu generierende Abschnitte:
{headings_text}

AUFGABE:
Generieren Sie für jeden oben genannten Abschnitt detaillierte, professionelle Inhalte.

WICHTIGE RAND bedingungen:
- Folgen Sie den Förderrichtlinien STRENG
- Erfinden Sie KEINE Daten - verwenden Sie NUR die bereitgestellten Firmeninformationen
- Folgen Sie dem Stil-Leitfaden STRENG
- Schreiben Sie AUSSCHLIESSLICH auf Deutsch
- Verwenden Sie NUR Absätze (keine Aufzählungspunkte)
- Fügen Sie KEINE Platzhalter, Fragen oder Haftungsausschlüsse ein
- Wenn Informationen unzureichend sind, generieren Sie plausible, professionelle Inhalte basierend auf dem verfügbaren Kontext

"""

    # Build complete prompt
    prompt = f"""Sie sind ein Expertenberater, der bei der Erstellung einer "Vorhabensbeschreibung" für einen Förderantrag hilft.

{rules_section}{company_section}{style_section}{task_section}

AUSGABEFORMAT:
Geben Sie NUR ein gültiges JSON-Objekt mit dieser exakten Struktur zurück:
{{
  "{section_ids[0] if section_ids else "section_id"}": "Generierter Absatztext...",
  "{section_ids[1] if len(section_ids) > 1 else "section_id"}": "Generierter Absatztext..."
}}

Die Schlüssel MÜSSEN exakt mit den Abschnitts-IDs aus der Liste oben übereinstimmen (z.B. "0", "1", "1.1", "2.3", etc.).
Die Werte müssen reiner deutscher Text in Absatzform sein.

Geben Sie KEIN Markdown-Format, KEINE Erklärungen und KEINEN Text außerhalb des JSON-Objekts zurück. Geben Sie NUR das JSON-Objekt zurück."""

    # Retry logic with JSON validation
    approx_tokens = len(prompt) // 4
    logger.info("LLM batch generation prompt size (chars): %s", len(prompt))
    logger.info("LLM batch generation prompt tokens: %s", approx_tokens)
    for attempt in range(max_retries + 1):
        try:
            response = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {
                        "role": "system",
                        "content": "Sie sind ein professioneller Berater, der sich auf Förderanträge spezialisiert hat. Sie schreiben klare, strukturierte und überzeugende Projektbeschreibungen auf Deutsch im formellen Fördermittel-Stil."
                    },
                    {
                        "role": "user",
                        "content": prompt
                    }
                ],
                temperature=0.7,
                response_format={"type": "json_object"},
                timeout=120.0  # 2 minute timeout for production safety
            )

            response_text = response.choices[0].message.content
            logger.info(f"OpenAI response received for batch (attempt {attempt + 1})")

            # Strict JSON validation
            try:
                generated_content = json.loads(response_text)

                # Validate that all expected section IDs are present
                missing_ids = [sid for sid in section_ids if sid not in generated_content]
                if missing_ids:
                    raise ValueError(f"Missing section IDs in response: {missing_ids}")

                # Validate that all values are strings
                for sid, content in generated_content.items():
                    if not isinstance(content, str):
                        raise ValueError(f"Content for section {sid} is not a string: {type(content)}")
                    if sid not in section_ids:
                        logger.warning(f"Unexpected section ID in response: {sid}")

                logger.info(f"Successfully validated JSON for batch with {len(generated_content)} sections")
                return generated_content

            except json.JSONDecodeError as e:
                error_msg = f"JSON parse error (attempt {attempt + 1}/{max_retries + 1}): {str(e)}. Response preview: {response_text[:200]}"
                logger.warning(error_msg)
                if attempt < max_retries:
                    continue
                raise ValueError(f"Failed to parse JSON after {max_retries + 1} attempts: {str(e)}") from e

            except ValueError as e:
                error_msg = f"JSON validation error (attempt {attempt + 1}/{max_retries + 1}): {str(e)}"
                logger.warning(error_msg)
                if attempt < max_retries:
                    continue
                raise

        except Exception as e:
            if attempt < max_retries:
                logger.warning(f"OpenAI API error (attempt {attempt + 1}/{max_retries + 1}): {str(e)}. Retrying...")
                continue
            logger.error(f"OpenAI API error after {max_retries + 1} attempts: {str(e)}")
            raise

    # Should never reach here, but just in case
    raise ValueError("Failed to generate content after all retries")


@router.post(
    "/documents/{document_id}/generate-content",
    response_model=DocumentResponse
)
def generate_content(
    document_id: int,
    db: Session = Depends(get_db),  # noqa: B008
    current_user: User = Depends(get_current_user)  # noqa: B008
):
    """
    ROLE: INITIAL CONTENT GENERATION

    Generate content for Vorhabensbeschreibung document using OpenAI with chunked generation.
    Requires company preprocessing to be completed.
    Generates content in batches of 3-5 sections for reliability and efficiency.

    This endpoint:
    - Creates initial content for empty sections
    - Uses _generate_batch_content() for generation logic
    - Assumes sections exist but have no content yet

    This endpoint must NOT:
    - Call _generate_section_content() (that's for editing only)
    - Be used for modifying existing content (use /chat instead)
    """

    # Load document
    document = _safe_get_document_by_id(document_id, db)
    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found"
        )

    # Verify document type
    if document.type != "vorhabensbeschreibung":
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Content generation only supported for vorhabensbeschreibung documents"
        )

    # Load associated company and verify ownership
    company = db.query(Company).filter(
        Company.id == document.company_id,
        Company.user_email == current_user.email
    ).first()
    if not company:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found"
        )

    # Check processing status
    if company.processing_status != "done":
        raise HTTPException(
            status_code=status.HTTP_409_CONFLICT,
            detail="Company preprocessing not finished"
        )

    # Load confirmed headings from document
    content_json = document.content_json
    if not content_json or "sections" not in content_json:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Document has no sections. Please create and confirm headings first."
        )

    sections = content_json["sections"]
    if not sections or len(sections) == 0:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Document has no confirmed headings. Please create and confirm headings first."
        )

    # Get OpenAI API key from environment
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="OpenAI API key not configured. Set OPENAI_API_KEY environment variable."
        )

    # Initialize OpenAI client
    try:
        client = OpenAI(api_key=api_key)
    except Exception as e:
        logger.error(f"Failed to initialize OpenAI client: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to initialize OpenAI client: {str(e)}"
        ) from e

    # Prepare company data (use cleaned versions)
    company_name = company.name or "Unknown Company"
    website_clean_text = company.website_clean_text or None
    transcript_clean = company.transcript_clean or None
    company_profile = company.company_profile  # PRIMARY factual source (structured JSON)

    # Get funding program rules if document has funding_program_id
    funding_program_rules = None
    if document.funding_program_id:
        from app.models import FundingProgramGuidelinesSummary
        summary = db.query(FundingProgramGuidelinesSummary).filter(
            FundingProgramGuidelinesSummary.funding_program_id == document.funding_program_id
        ).first()
        if summary:
            funding_program_rules = summary.rules_json
            logger.info(f"Using funding program rules for funding_program_id={document.funding_program_id}")

    # Get style profile (system-level, from AlteVorhabensbeschreibung)
    style_profile = None
    from app.models import AlteVorhabensbeschreibungStyleProfile
    style_profile_record = db.query(AlteVorhabensbeschreibungStyleProfile).first()
    if style_profile_record:
        style_profile = style_profile_record.style_summary_json
        logger.info(f"Using style profile (hash: {style_profile_record.combined_hash[:10]}...)")
    else:
        logger.warning("No style profile found - generation will use default style guidelines")

    # Filter out milestone tables from content generation (they should not be AI-generated)
    text_sections = [s for s in sections if s.get("type") != "milestone_table"]
    milestone_sections = [s for s in sections if s.get("type") == "milestone_table"]

    # Split only text sections into batches (3-5 sections per batch)
    batches = _split_sections_into_batches(text_sections, batch_size=4)
    logger.info(f"Split {len(text_sections)} text sections into {len(batches)} batches for document {document_id}")
    if milestone_sections:
        logger.info(f"Excluded {len(milestone_sections)} milestone table(s) from content generation")

    # Initialize section content map (preserve existing content for all sections)
    section_content_map = {}
    for section in sections:
        section_id = section.get("id", "")
        existing_content = section.get("content", "")
        section_content_map[section_id] = existing_content

    # Process each batch
    successful_batches = 0
    failed_batches = []

    for batch_idx, batch in enumerate(batches):
        try:
            logger.info(f"Processing batch {batch_idx + 1}/{len(batches)} with {len(batch)} sections")

            # Generate content for this batch
            # NOTE: This calls _generate_batch_content (INITIAL GENERATION role)
            # This is correct - we are generating initial content, not editing existing content
            batch_content = _generate_batch_content(
                client=client,
                batch_sections=batch,
                company_name=company_name,
                company_profile=company_profile,  # PRIMARY factual source
                website_clean_text=website_clean_text,  # Contextual enrichment
                transcript_clean=transcript_clean,  # Contextual enrichment
                company_id=company.id,  # Guardrail A: Pass company_id for logging
                funding_program_rules=funding_program_rules,  # Rules and guidelines
                style_profile=style_profile,  # Style guide
                max_retries=2
            )

            # Merge batch content into section map
            for section_id, content in batch_content.items():
                if section_id in section_content_map:
                    section_content_map[section_id] = content
                else:
                    logger.warning(f"Generated content for unexpected section ID: {section_id}")

            # Persist incrementally after each successful batch
            updated_sections = []
            for section in sections:
                section_id = section.get("id", "")
                section_title = section.get("title", "")
                section_type = section.get("type", "text")  # Preserve type field
                content = section_content_map.get(section_id, section.get("content", ""))

                # Don't overwrite milestone table content with text - skip generation for milestone tables
                if section_type == "milestone_table":
                    # Keep existing milestone table structure, don't replace with generated text
                    content = section.get("content", "")

                updated_sections.append({
                    "id": section_id,
                    "title": section_title,
                    "type": section_type,  # Preserve type field
                    "content": content
                })

            document.content_json = {"sections": updated_sections}
            db.commit()
            db.refresh(document)

            successful_batches += 1
            logger.info(f"Successfully processed and persisted batch {batch_idx + 1}/{len(batches)}")

        except Exception as e:
            # Log error but continue with other batches
            batch_section_ids = [s.get("id", "") for s in batch]
            error_msg = f"Failed to generate content for batch {batch_idx + 1} (sections: {batch_section_ids}): {str(e)}"
            logger.error(error_msg)
            logger.error(f"Full traceback:\n{traceback.format_exc()}")
            failed_batches.append({
                "batch_index": batch_idx + 1,
                "section_ids": batch_section_ids,
                "error": str(e)
            })
            # Continue with next batch - partial success is preserved

    # Final status check
    if successful_batches == 0:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to generate content for all batches. Errors: {[b['error'] for b in failed_batches]}"
        )

    if failed_batches:
        logger.warning(f"Completed generation with {len(failed_batches)} failed batches out of {len(batches)} total")
        # Partial success - return what we have, but log the failures

    # Final refresh to ensure we return the latest state
    db.refresh(document)
    logger.info(f"Successfully completed content generation for document {document_id}: {successful_batches}/{len(batches)} batches succeeded")

    return document


def _normalize_section_id(section_id: str) -> str:
    """
    Normalize section ID by:
    1. Converting commas to dots (e.g., "1,1" -> "1.1")
    2. Stripping trailing dots and whitespace
    3. Removing extra spaces

    Examples:
    - "2.1." -> "2.1"
    - "2.1 " -> "2.1"
    - "1,1" -> "1.1"
    - "1, 1" -> "1.1"
    - "2.1.2" -> "2.1.2"
    """
    if not section_id:
        return section_id

    # Convert commas to dots (common mistake: "1,1" instead of "1.1")
    normalized = section_id.replace(',', '.')

    # Remove spaces around dots (e.g., "1. 1" -> "1.1", "1 , 1" -> "1.1")
    normalized = re.sub(r'\s*\.\s*', '.', normalized)

    # Strip trailing dots and whitespace
    normalized = normalized.rstrip('.').strip()

    return normalized


def _find_section_by_title(
    user_input: str,
    sections: List[dict],
    threshold: float = 0.8
) -> Optional[str]:
    """
    Find section ID by matching title using fuzzy matching.

    Strategy:
    1. Exact match (normalized, case-insensitive)
    2. Partial match (title contains input or vice versa)
    3. Fuzzy match using SequenceMatcher (similarity >= threshold)

    Returns section_id if found, None otherwise.
    """
    if not user_input or not sections:
        return None

    from difflib import SequenceMatcher

    # Normalize user input
    user_input_normalized = user_input.lower().strip()

    # Build title-to-ID mapping with normalized titles
    title_matches = []  # List of (section_id, normalized_title, similarity_score)

    for section in sections:
        section_id = section.get("id", "")
        section_title = section.get("title", "")

        if not section_title or not section_id:
            continue

        # Remove numbering prefix (e.g., "2.1. Firmengeschichte" -> "Firmengeschichte")
        clean_title = re.sub(r'^[\d.]+\.\s*', '', section_title).strip()
        normalized_title = clean_title.lower()

        # Strategy 1: Exact match (normalized)
        if normalized_title == user_input_normalized:
            logger.debug(f"Exact title match: '{user_input}' -> section {section_id}")
            return section_id

        # Strategy 2: Partial match (contains)
        if user_input_normalized in normalized_title or normalized_title in user_input_normalized:
            # Calculate similarity for ranking
            similarity = SequenceMatcher(None, user_input_normalized, normalized_title).ratio()
            title_matches.append((section_id, clean_title, similarity))
            logger.debug(f"Partial title match: '{user_input}' ~ '{clean_title}' (similarity: {similarity:.2f})")
            continue

        # Strategy 3: Fuzzy match
        similarity = SequenceMatcher(None, user_input_normalized, normalized_title).ratio()
        if similarity >= threshold:
            title_matches.append((section_id, clean_title, similarity))
            logger.debug(f"Fuzzy title match: '{user_input}' ~ '{clean_title}' (similarity: {similarity:.2f})")

    # If we have matches, return the best one (highest similarity)
    if title_matches:
        # Sort by similarity (descending), then by section_id for consistency
        title_matches.sort(key=lambda x: (-x[2], x[0]))
        best_match = title_matches[0]
        logger.info(f"Best title match: '{user_input}' -> section {best_match[0]} (similarity: {best_match[2]:.2f})")
        return best_match[0]

    return None


def _parse_section_changes_enhanced(user_message: str, valid_section_ids: List[str], sections: List[dict] = None) -> List[dict]:
    """
    Enhanced flexible parser that understands various natural language formats.
    This parser is more permissive than the original but still deterministic and safe.

    Supports formats like:
    - "Section 2.1: make it more concise"
    - "2.1: make it innovative"
    - "2.1 innovation"
    - "2.1 - innovation"
    - "Update 2.3 to emphasize sustainability"
    - "2.1: concise. 2.2: more technical"
    - "2.1, 2.3, and 2.5: make them all innovative"

    Returns empty list if nothing reliable is found (no guessing).
    """
    logger.debug(f"_parse_section_changes_enhanced called with message: '{user_message}', valid_section_ids: {valid_section_ids}")
    changes = []
    message = user_message.strip()

    # Strategy: Find all section references first, then extract instructions for each

    # Find all potential section references with their positions
    section_matches = []

    # NEW: Try to find sections by title first (if sections provided)
    if sections:
        # Pattern 1: "TitleName: instruction" or "TitleName - instruction"
        title_patterns = [
            re.compile(r'([A-Za-zÄÖÜäöüß][A-Za-zÄÖÜäöüß\s]{2,40}?)\s*[:]\s*(.+?)(?=\n|$|[\d.]+\s*:|[A-Za-zÄÖÜäöüß][A-Za-zÄÖÜäöüß\s]{2,40}?\s*:)', re.IGNORECASE),
            re.compile(r'([A-Za-zÄÖÜäöüß][A-Za-zÄÖÜäöüß\s]{2,40}?)\s*[-]\s*(.+?)(?=\n|$|[\d.]+\s*[-:]|section|abschnitt)', re.IGNORECASE),
        ]

        for pattern in title_patterns:
            for match in pattern.finditer(message):
                potential_title = match.group(1).strip()
                instruction = match.group(2).strip() if len(match.groups()) > 1 else ""

                # Try to find section by title
                section_id = _find_section_by_title(potential_title, sections, threshold=0.8)
                if section_id and section_id in valid_section_ids and instruction and len(instruction) > 2:
                    changes.append({
                        "section_id": section_id,
                        "instruction": instruction
                    })
                    logger.info(f"Found section by title: '{potential_title}' -> {section_id}, instruction: '{instruction}'")
                    # Remove this part from message to avoid duplicate parsing
                    message = message[:match.start()] + message[match.end():]
                    break
            if changes:
                break  # If we found a title match, don't try other patterns

        # Pattern 2: Standalone title word (e.g., just "Firmengeschichte" followed by instruction)
        # Only try this if no other patterns matched
        if not changes:
            # Look for capitalized words that might be section titles
            standalone_title_pattern = re.compile(r'^([A-ZÄÖÜ][a-zäöüß]+(?:\s+[A-ZÄÖÜ][a-zäöüß]+){0,3})\s+(.+?)(?=\n|$|section|abschnitt|[\d.]+\s*[:])', re.MULTILINE | re.IGNORECASE)
            for match in standalone_title_pattern.finditer(message):
                potential_title = match.group(1).strip()
                instruction = match.group(2).strip() if len(match.groups()) > 1 else ""

                # Skip if it looks like a section number pattern (with dots or commas)
                if re.match(r'^[\d.,]+$', potential_title):
                    continue

                # Try to find section by title
                section_id = _find_section_by_title(potential_title, sections, threshold=0.8)
                if section_id and section_id in valid_section_ids and instruction and len(instruction) > 2:
                    changes.append({
                        "section_id": section_id,
                        "instruction": instruction
                    })
                    logger.info(f"Found section by standalone title: '{potential_title}' -> {section_id}, instruction: '{instruction}'")
                    break

    # Strategy: Find all section references first, then extract instructions for each

    # Find all potential section references with their positions
    section_matches = []

    # Pattern 1: "Section X.Y" or "Abschnitt X.Y" (also matches commas: "Section 1,1")
    pattern1 = re.compile(r'(?:section|abschnitt)\s+([\d.,]+)', re.IGNORECASE)
    for match in pattern1.finditer(message):
        section_id = _normalize_section_id(match.group(1))
        if section_id in valid_section_ids:
            section_matches.append({
                'id': section_id,
                'start': match.start(),
                'end': match.end(),
                'type': 'explicit'
            })

    # Pattern 2: "X.Y:" (direct section ID with colon, also matches commas: "1,1:")
    # Use negative lookbehind to avoid matching partial IDs (e.g., "4" from "2.4")
    pattern2 = re.compile(r'(?<![.,\d])([\d.,]+)\s*:', re.MULTILINE)
    for match in pattern2.finditer(message):
        section_id = _normalize_section_id(match.group(1))
        if section_id in valid_section_ids:
            # Avoid duplicates from pattern1
            if not any(m['id'] == section_id and m['start'] == match.start() for m in section_matches):
                section_matches.append({
                    'id': section_id,
                    'start': match.start(),
                    'end': match.end(),
                    'type': 'colon'
                })

    # Pattern 3: "X.Y -" or "X.Y-" (dash format, also matches commas: "1,1 -")
    # Match section ID followed by dash, ensuring we get the full ID (e.g., "2.4" not just "4")
    # Use negative lookbehind to avoid matching partial IDs
    pattern3 = re.compile(r'(?<![.,\d])([\d.,]+)\s*-\s*', re.MULTILINE)
    for match in pattern3.finditer(message):
        section_id = _normalize_section_id(match.group(1))
        if section_id in valid_section_ids:
            # Avoid duplicates
            if not any(m['id'] == section_id and abs(m['start'] - match.start()) < 5 for m in section_matches):
                section_matches.append({
                    'id': section_id,
                    'start': match.start(),
                    'end': match.end(),
                    'type': 'dash'
                })

    # Pattern 4: "Update/Rewrite section X.Y" or action verbs with section (also matches commas)
    pattern4 = re.compile(
        r'(?:update|rewrite|change|modify|edit|überarbeite|aktualisiere|ändere|verbessere|erweitere|kürze|betone)\s+(?:section|abschnitt)?\s*([\d.,]+)',
        re.IGNORECASE
    )
    for match in pattern4.finditer(message):
        section_id = _normalize_section_id(match.group(1))
        if section_id in valid_section_ids:
            # Avoid duplicates
            if not any(m['id'] == section_id and abs(m['start'] - match.start()) < 10 for m in section_matches):
                section_matches.append({
                    'id': section_id,
                    'start': match.start(),
                    'end': match.end(),
                    'type': 'action'
                })

    # Pattern 5: Standalone section ID at start of line or after punctuation (also matches commas)
    pattern5 = re.compile(r'(?:^|[\n\.])\s*([\d.,]+)\s+(?![\d.,])', re.MULTILINE)
    for match in pattern5.finditer(message):
        section_id = _normalize_section_id(match.group(1))
        if section_id in valid_section_ids:
            # Only add if it's clearly a section reference (not part of a number)
            # Check if followed by meaningful text (not just another number)
            pos = match.end()
            if pos < len(message):
                next_chars = message[pos:pos+20].strip()
                # If followed by action words or meaningful text, it's likely a section reference
                if next_chars and not re.match(r'^[\d.\s,]+$', next_chars):
                    # Avoid duplicates
                    if not any(m['id'] == section_id and abs(m['start'] - match.start()) < 5 for m in section_matches):
                        section_matches.append({
                            'id': section_id,
                            'start': match.start(),
                            'end': match.end(),
                            'type': 'standalone'
                        })

    # Remove duplicates (keep first occurrence)
    seen_ids = set()
    unique_matches = []
    for match in section_matches:
        if match['id'] not in seen_ids:
            seen_ids.add(match['id'])
            unique_matches.append(match)
            logger.debug(f"Added section match: id={match['id']}, type={match['type']}, position={match['start']}")
        else:
            logger.debug(f"Skipped duplicate section match: id={match['id']}, type={match['type']}, position={match['start']}")

    # Sort by position in message
    unique_matches.sort(key=lambda x: x['start'])
    logger.debug(f"Final unique matches after sorting: {[m['id'] for m in unique_matches]}")

    # Extract instruction for each section
    for i, sec_match in enumerate(unique_matches):
        section_id = sec_match['id']
        instruction_start = sec_match['end']

        # Find where this instruction ends (next section or end of message)
        if i + 1 < len(unique_matches):
            instruction_end = unique_matches[i + 1]['start']
        else:
            instruction_end = len(message)

        # Extract instruction text
        instruction_text = message[instruction_start:instruction_end].strip()

        # Clean up instruction
        # Remove leading separators (colon, dash, whitespace)
        # Note: dash must be escaped or at end of character class to avoid being interpreted as range
        instruction_text = re.sub(r'^[-:\s]+', '', instruction_text)

        # Remove trailing separators
        instruction_text = re.sub(r'\s*[-:\s]*$', '', instruction_text)

        # Remove trailing punctuation that might be from sentence structure
        instruction_text = re.sub(r'[.,;]+$', '', instruction_text).strip()

        # Validate instruction is meaningful
        if instruction_text and len(instruction_text) > 2:
            # Check if it's just another section reference (skip if so)
            # Fix: dash must be at beginning or end of character class
            # Also check for commas in section references
            if not re.match(r'^[\d.,]+\s*[-:\s]', instruction_text):
                logger.debug(f"Found valid change: section_id={section_id}, instruction='{instruction_text}'")
                changes.append({
                    "section_id": section_id,
                    "instruction": instruction_text
                })
            else:
                logger.debug(f"Skipping instruction that looks like section reference: '{instruction_text}'")
        else:
            logger.debug(f"Skipping instruction (too short or empty): '{instruction_text}'")

    logger.debug(f"_parse_section_changes_enhanced returning {len(changes)} changes: {changes}")
    return changes


def _parse_section_changes(user_message: str, valid_section_ids: List[str], sections: List[dict] = None) -> List[dict]:
    """
    Parse user message to extract section IDs and their corresponding instructions.
    Returns a list of {section_id, instruction} dictionaries.

    This is a deterministic, rule-based parser (not LLM-based).
    Supports multiple formats:
    - "Section 2.1: make it more concise"
    - "Rewrite section 2.1 to emphasize sustainability"
    - "Section 1.1: make more concise. Section 2.3: emphasize innovation"
    - "2.1: make it shorter"
    """
    changes = []

    # Normalize message
    message = user_message.strip()

    # NEW: Try to find sections by title first (if sections provided)
    if sections:
        # Pattern 1: "TitleName: instruction" or "TitleName - instruction"
        title_patterns = [
            re.compile(r'([A-Za-zÄÖÜäöüß][A-Za-zÄÖÜäöüß\s]{2,40}?)\s*[:]\s*(.+?)(?=\n|$|[\d.]+\s*:|section|abschnitt)', re.IGNORECASE),
            re.compile(r'([A-Za-zÄÖÜäöüß][A-Za-zÄÖÜäöüß\s]{2,40}?)\s*[-]\s*(.+?)(?=\n|$|[\d.]+\s*[-:]|section|abschnitt)', re.IGNORECASE),
        ]

        for pattern in title_patterns:
            for match in pattern.finditer(message):
                potential_title = match.group(1).strip()
                instruction = match.group(2).strip() if len(match.groups()) > 1 else ""

                # Try to find section by title
                section_id = _find_section_by_title(potential_title, sections, threshold=0.8)
                if section_id and section_id in valid_section_ids and instruction and len(instruction) > 2:
                    changes.append({
                        "section_id": section_id,
                        "instruction": instruction
                    })
                    logger.info(f"Found section by title: '{potential_title}' -> {section_id}, instruction: '{instruction}'")
                    message = message[:match.start()] + message[match.end():]
                    break
            if changes:
                break

        # Pattern 2: Standalone title word (e.g., just "Firmengeschichte" followed by instruction)
        # Only try this if no other patterns matched
        if not changes:
            # Look for capitalized words that might be section titles
            standalone_title_pattern = re.compile(r'^([A-ZÄÖÜ][a-zäöüß]+(?:\s+[A-ZÄÖÜ][a-zäöüß]+){0,3})\s+(.+?)(?=\n|$|section|abschnitt|[\d.]+\s*[:])', re.MULTILINE | re.IGNORECASE)
            for match in standalone_title_pattern.finditer(message):
                potential_title = match.group(1).strip()
                instruction = match.group(2).strip() if len(match.groups()) > 1 else ""

                # Skip if it looks like a section number pattern (with dots or commas)
                if re.match(r'^[\d.,]+$', potential_title):
                    continue

                # Try to find section by title
                section_id = _find_section_by_title(potential_title, sections, threshold=0.8)
                if section_id and section_id in valid_section_ids and instruction and len(instruction) > 2:
                    changes.append({
                        "section_id": section_id,
                        "instruction": instruction
                    })
                    logger.info(f"Found section by standalone title: '{potential_title}' -> {section_id}, instruction: '{instruction}'")
                    break

    # Pattern 1: "Section X.Y: instruction" or "Abschnitt X.Y: instruction" (with colon, also matches commas)
    pattern1 = re.compile(r'(?:section|abschnitt)\s+([\d.,]+)\s*:+\s*(.+?)(?=(?:section|abschnitt)\s+[\d.,]+|$)', re.IGNORECASE | re.DOTALL)
    matches1 = pattern1.findall(message)
    for section_id, instruction in matches1:
        section_id = _normalize_section_id(section_id)
        instruction = instruction.strip()
        if section_id in valid_section_ids and instruction and len(instruction) > 3:
            changes.append({"section_id": section_id, "instruction": instruction})

    # Pattern 2: "X.Y: instruction" (direct section ID with colon, also matches commas: "1,1:")
    pattern2 = re.compile(r'^([\d.,]+)\s*:+\s*(.+?)(?=\n|$|[\d.,]+\s*:+)', re.MULTILINE | re.IGNORECASE | re.DOTALL)
    matches2 = pattern2.findall(message)
    for section_id, instruction in matches2:
        section_id = _normalize_section_id(section_id)
        instruction = instruction.strip()
        if section_id in valid_section_ids and instruction and len(instruction) > 3:
            # Avoid duplicates
            if not any(c["section_id"] == section_id for c in changes):
                changes.append({"section_id": section_id, "instruction": instruction})

    # Pattern 3: "Rewrite/Update section X.Y to..." or "Überarbeite Abschnitt X.Y zu..." (also matches commas)
    pattern3 = re.compile(r'(?:rewrite|update|change|modify|edit|überarbeite|aktualisiere|ändere|verbessere|erweitere|kürze|betone)\s+(?:section|abschnitt)?\s*([\d.,]+)\s+(?:to|zu|mit|dass|damit|so dass|um)\s+(.+)', re.IGNORECASE | re.DOTALL)
    matches3 = pattern3.findall(message)
    for section_id, instruction in matches3:
        section_id = _normalize_section_id(section_id)
        instruction = instruction.strip()
        if section_id in valid_section_ids and instruction and len(instruction) > 3:
            # Avoid duplicates
            if not any(c["section_id"] == section_id for c in changes):
                changes.append({"section_id": section_id, "instruction": instruction})

    # Pattern 4: "Section X.Y" followed by instruction (without colon, separated by newline or period, also matches commas)
    # Only if no other patterns matched
    if not changes:
        pattern4 = re.compile(r'(?:section|abschnitt)\s+([\d.,]+)\s+([^\n\.]+)', re.IGNORECASE)
        matches4 = pattern4.findall(message)
        for section_id, instruction in matches4:
            section_id = _normalize_section_id(section_id)
            instruction = instruction.strip()
            # Only add if it looks like an instruction (not just "section X.Y" alone)
            if section_id in valid_section_ids and instruction and len(instruction) > 5:
                changes.append({"section_id": section_id, "instruction": instruction})

    # Remove duplicates (keep first occurrence)
    seen = set()
    unique_changes = []
    for change in changes:
        if change["section_id"] not in seen:
            seen.add(change["section_id"])
            unique_changes.append(change)

    return unique_changes


def _validate_section_changes(changes: List[dict], valid_section_ids: List[str]) -> Tuple[bool, Optional[str]]:
    """
    Validate that all changes have valid section IDs and instructions.
    Returns (is_valid, error_message).
    """
    if not changes:
        return False, None

    # Check all section IDs are valid
    invalid_ids = [c["section_id"] for c in changes if c["section_id"] not in valid_section_ids]
    if invalid_ids:
        return False, f"Ungültige Abschnitts-IDs gefunden: {', '.join(invalid_ids)}. Bitte geben Sie gültige Abschnittsnummern an (z.B. 1.1, 2.3)."

    # Check all have instructions
    missing_instructions = [c["section_id"] for c in changes if not c.get("instruction") or len(c["instruction"].strip()) < 3]
    if missing_instructions:
        return False, f"Bitte geben Sie für Abschnitt {missing_instructions[0]} eine konkrete Anweisung an, was geändert werden soll."

    return True, None


def _determine_clarification_needed(
    user_message: str,
    valid_section_ids: List[str],
    last_edited_sections: Optional[List[str]] = None
) -> Optional[str]:
    """
    Determine if clarification is needed before calling LLM.
    Returns conversational clarification question if needed, None if ready to proceed.
    This is a deterministic, rule-based check (not LLM-based).

    Uses context (last_edited_sections) to suggest sections but never auto-applies.
    """
    logger.debug(f"_determine_clarification_needed called with message: '{user_message}', last_edited_sections: {last_edited_sections}")
    # Try enhanced parser first
    try:
        changes_enhanced = _parse_section_changes_enhanced(user_message, valid_section_ids)
        logger.debug(f"Enhanced parser returned {len(changes_enhanced)} changes")
    except Exception as e:
        logger.error(f"Error in enhanced parser: {str(e)}", exc_info=True)
        changes_enhanced = []
    if changes_enhanced:
        is_valid, error_msg = _validate_section_changes(changes_enhanced, valid_section_ids)
        if is_valid:
            return None  # No clarification needed
        if error_msg:
            return error_msg  # Return validation error

    # Fallback to original parser
    try:
        changes_original = _parse_section_changes(user_message, valid_section_ids)
        logger.debug(f"Original parser returned {len(changes_original)} changes")
        if changes_original:
            is_valid, error_msg = _validate_section_changes(changes_original, valid_section_ids)
            if is_valid:
                logger.debug("Original parser found valid changes, no clarification needed")
                return None  # No clarification needed
            if error_msg:
                logger.debug(f"Original parser found changes but validation failed: {error_msg}")
                return error_msg  # Return validation error
    except Exception as e:
        logger.error(f"Error in original parser: {str(e)}", exc_info=True)

    # No valid changes found - need clarification
    # Check if message has action verbs (user wants to do something)
    action_pattern = re.compile(
        r'(?:make|update|change|edit|improve|fix|add|remove|rewrite|'
        r'überarbeite|aktualisiere|ändere|verbessere|erweitere|kürze|betone|'
        r'innovative|innovativ|shorter|longer|concise|detailed|technical|technisch)',
        re.IGNORECASE
    )
    has_action = bool(action_pattern.search(user_message))

    # Check if any section IDs mentioned (even if not parsed correctly)
    section_refs = re.findall(r'\b([\d.]+)\b', user_message)
    potential_sections = [s for s in section_refs if s in valid_section_ids]
    invalid_sections = [s for s in section_refs if s not in valid_section_ids and re.match(r'^\d+(\.\d+)*$', s)]

    # Case 1: Invalid section IDs found
    if invalid_sections:
        unique_invalid = list(set(invalid_sections))
        if len(unique_invalid) == 1:
            return f"Ich konnte Abschnitt {unique_invalid[0]} nicht finden. Bitte geben Sie eine gültige Abschnittsnummer an (z.B. 2.1, 3.2)."
        return f"Ich konnte die Abschnittsnummern {', '.join(unique_invalid)} nicht finden. Bitte geben Sie gültige Abschnittsnummern an (z.B. 2.1, 3.2)."

    # Case 2: No sections mentioned at all
    if not potential_sections:
        if has_action:
            # User wants to do something but didn't specify section
            if last_edited_sections and len(last_edited_sections) > 0:
                if len(last_edited_sections) == 1:
                    return f"Meinen Sie Abschnitt {last_edited_sections[0]}? Bitte bestätigen Sie, oder geben Sie die Abschnittsnummer an (z.B. 2.1 oder 2.1 und 2.3)."
                else:
                    sections_str = ", ".join(last_edited_sections)
                    return "Welche Abschnitte sollen aktualisiert werden? Sie können mehrere angeben (z.B. 2.1 oder 2.1 und 2.3)."
            return "Welche Abschnitte sollen aktualisiert werden? Bitte geben Sie Abschnittsnummern an (z.B. 2.1 oder 2.1 und 2.3)."
        return "Bitte geben Sie an, welche Abschnitte geändert werden sollen und was genau geändert werden soll (z.B. '2.1: make it innovative' oder 'Section 2.1: make it more concise')."

    # Case 3: Sections mentioned but couldn't parse instruction
    if has_action:
        if len(potential_sections) == 1:
            return f"Was soll in Abschnitt {potential_sections[0]} geändert werden? Bitte geben Sie eine klarere Anweisung an (z.B. 'make it more innovative' oder 'fix the style')."
        else:
            sections_str = ", ".join(potential_sections)
            return f"Was soll in den Abschnitten {sections_str} geändert werden? Bitte geben Sie für jeden Abschnitt eine Anweisung an (z.B. '2.1: make it innovative. 2.2: fix the style')."

    # Case 4: Sections mentioned but no action verb
    if len(potential_sections) == 1:
        return f"Was soll in Abschnitt {potential_sections[0]} geändert werden? Bitte geben Sie eine Anweisung an (z.B. 'make it more innovative' oder 'make it shorter')."
    else:
        sections_str = ", ".join(potential_sections)
        return f"Was soll in den Abschnitten {sections_str} geändert werden? Bitte geben Sie für jeden Abschnitt eine Anweisung an."

    # Fallback (should not reach here)
    return "Bitte geben Sie an, welche Abschnitte geändert werden sollen und was genau geändert werden soll."


def _generate_section_content(
    client: OpenAI,
    section_id: str,
    section_title: str,
    current_content: str,
    instruction: str,
    company_name: str,
    company_profile: Optional[dict] = None,
    website_clean_text: Optional[str] = None,
    transcript_clean: Optional[str] = None,
    company_id: Optional[int] = None,
    style_profile: Optional[Dict[str, Any]] = None
) -> str:
    """
    ROLE: SECTION EDITOR

    Modifies EXISTING section content based on user editing instructions.
    Used ONLY for chat-based editing via /chat endpoint.

    This function:
    - Assumes sections already have content that needs modification
    - Focuses on targeted editing and refinement
    - Is constrained and conservative (preserves existing structure)
    - Uses existing content as the primary basis for changes

    This function must NOT:
    - Be used for initial content generation
    - Regenerate sections from scratch
    - Be called from /generate-content endpoint

    Parameters:
    - current_content: The existing section content that will be modified
    - instruction: User's editing instruction (e.g., "make it more concise")

    Returns the updated section content as a string.
    """
    # Remove numbering prefix from title
    clean_title = re.sub(r'^[\d.]+\.\s*', '', section_title)

    # IMPORTANT:
    # This prompt is for EDITING existing content only.
    # Do NOT reuse this prompt for initial content generation.
    # For initial generation, use _generate_batch_content() instead.
    # This prompt assumes existing content exists and must be modified, not created.

    # Format company context using cleaned data
    company_context = _format_company_context_for_prompt(
        company_profile=company_profile,
        company_name=company_name,
        website_clean_text=website_clean_text,
        transcript_clean=transcript_clean,
        company_id=company_id
    )

    # Build style guide section from style profile
    style_guide = ""
    if style_profile:
        style_parts = []
        
        if style_profile.get("structure_patterns"):
            patterns = style_profile["structure_patterns"]
            if isinstance(patterns, list) and patterns:
                style_parts.append("Strukturmuster:\n" + "\n".join(f"- {p}" for p in patterns))
        
        if style_profile.get("tone_characteristics"):
            tone = style_profile["tone_characteristics"]
            if isinstance(tone, list) and tone:
                style_parts.append("Ton und Charakteristik:\n" + "\n".join(f"- {t}" for t in tone))
        
        if style_profile.get("writing_style_rules"):
            rules = style_profile["writing_style_rules"]
            if isinstance(rules, list) and rules:
                style_parts.append("Schreibstil-Regeln:\n" + "\n".join(f"- {r}" for r in rules))
        
        if style_profile.get("storytelling_flow"):
            flow = style_profile["storytelling_flow"]
            if isinstance(flow, list) and flow:
                style_parts.append("Erzählstruktur:\n" + "\n".join(f"- {f}" for f in flow))
        
        if style_parts:
            style_guide = "=== STIL-LEITFADEN ===\n\n" + "\n\n".join(style_parts) + "\n\n"
            style_guide += "WICHTIG: Folgen Sie diesen Stilrichtlinien bei der Überarbeitung.\n\n"
    else:
        style_guide = "=== STIL-LEITFADEN ===\n\n"
        style_guide += "- Verwenden Sie formelle Fördermittel-/Geschäftssprache\n"
        style_guide += "- Professioneller, überzeugender Ton\n"
        style_guide += "- Klare Absatzstruktur\n\n"

    # Build prompt with style guide
    instruction_text = instruction or ""
    prompt = f"""{style_guide}SIE SIND EIN REDAKTEUR, KEIN AUTOR.

- Der folgende Abschnitt EXISTIERT bereits.
- Ihre Aufgabe ist es, den bestehenden Text gezielt zu überarbeiten.
- Ersetzen Sie NICHT den gesamten Inhalt, außer die Benutzeranweisung verlangt dies ausdrücklich.
- Bewahren Sie Struktur, Kernaussagen und Tonalität des bestehenden Textes.

PRIMÄRE GRUNDLAGE:

- Der bestehende Abschnittstext ist die wichtigste Grundlage.
- Änderungen müssen sich auf den vorhandenen Inhalt beziehen.
- Fügen Sie neue Informationen nur hinzu, wenn sie logisch an den bestehenden Text anschließen.

Aktueller Abschnitt:
- Abschnitts-ID: {section_id}
- Titel: {clean_title}
- Aktueller Inhalt: {current_content}

Benutzeranweisung: <user_instruction>
{instruction_text}
</user_instruction>

KONTEXTNUTZUNG:

- Verwenden Sie Firmeninformationen ausschließlich zur Präzisierung oder inhaltlichen Stützung.
- Fügen Sie keine neuen Themen ein, die im bestehenden Abschnitt nicht bereits angelegt sind.
- Vermeiden Sie generische Aussagen ohne Bezug zum aktuellen Abschnitt.

Firmeninformationen (NUR ZUR STÜTZUNG):
{company_context}

UMGANG MIT ALLGEMEINEN ANWEISUNGEN:

- Bei unspezifischen Anweisungen wie „Inhalt hinzufügen", „verbessern" oder „ausbauen":
  - Erweitern Sie den bestehenden Text moderat (ca. +20–40%).
  - Vertiefen Sie bestehende Aussagen, anstatt neue Themen zu eröffnen.

- Bei spezifischen Anweisungen wie „kürzer", „präziser" oder „technischer":
  - Passen Sie den Text entsprechend an, behalten Sie aber die Kernaussagen bei.

- Bei Anweisungen wie „rewrite" oder „komplett neu":
  - Formulieren Sie den Text neu, aber behalten Sie die inhaltlichen Kernpunkte bei.
  - Erweitern Sie moderat (ca. +30–50%), nicht exzessiv.

ABSCHNITTSFOKUS:

- Der überarbeitete Text muss inhaltlich eindeutig zum Titel des Abschnitts passen.
- Fügen Sie keine Themen hinzu, die zu anderen Abschnitten gehören.
- Ändern Sie NICHT den Titel oder die Struktur des Abschnitts.

STIL UND SPRACHE:

- Schreiben Sie ausschließlich auf Deutsch.
- Verwenden Sie einen sachlichen, formellen Fördermittel-Stil.
- Schreiben Sie in zusammenhängenden Absätzen (keine Aufzählungen).
- Keine Meta-Kommentare, keine Hinweise auf KI, keine Platzhalter.
- Stellen Sie KEINE Fragen.
- Fügen Sie KEINE Zitate oder Haftungsausschlüsse ein.
- Erwähnen Sie KEINE vorherigen Versionen oder Änderungen.

WICHTIG:

- Ändern Sie NICHT den Abschnittstitel.
- Fügen Sie KEINE neuen Abschnitte hinzu.
- Der Inhalt muss mit den Firmeninformationen übereinstimmen.
- Geben Sie NUR den überarbeiteten Absatztext zurück (kein JSON, kein Markdown, keine Erklärungen)."""

    approx_tokens = len(prompt) // 4
    logger.info("LLM section edit prompt size (chars): %s", len(prompt))
    logger.info("LLM section edit prompt tokens: %s", approx_tokens)
    try:
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {
                    "role": "system",
                    "content": "Sie sind ein professioneller Redakteur, der bestehende Abschnitte von Vorhabensbeschreibungen gezielt überarbeitet. Sie sind KEIN Autor, der Inhalte neu erstellt. Ihre Aufgabe ist die präzise Bearbeitung vorhandener Texte auf Deutsch im formellen Fördermittel-Stil."
                },
                {
                    "role": "user",
                    
                    "content": prompt
                }
            ],
            temperature=0.7,
            max_tokens=2000,
            timeout=120.0  # 2 minute timeout for production safety
        )

        generated_content = response.choices[0].message.content.strip()

        # Clean up any markdown or JSON artifacts
        generated_content = re.sub(r'^```(?:json|markdown)?\s*\n?', '', generated_content)
        generated_content = re.sub(r'\n?```\s*$', '', generated_content)
        generated_content = generated_content.strip()

        return generated_content

    except Exception as e:
        logger.error(f"Failed to generate content for section {section_id}: {str(e)}")
        raise


def _is_question(message: str) -> bool:
    """
    Detect if a message is a question based on patterns.
    Questions typically start with question words or end with '?'.
    """
    message_lower = message.lower().strip()

    # Check for question mark
    if message_lower.endswith('?'):
        return True

    # Check for question words at the start
    question_words = ['what', 'how', 'why', 'when', 'where', 'who', 'which', 'can', 'could', 'should', 'would', 'is', 'are', 'was', 'were', 'do', 'does', 'did', 'will', 'tell me', 'explain', 'describe']
    first_word = message_lower.split()[0] if message_lower.split() else ""

    if first_word in question_words:
        return True

    # Check for question patterns
    question_patterns = [
        r'^what\s+',
        r'^how\s+',
        r'^why\s+',
        r'^when\s+',
        r'^where\s+',
        r'^who\s+',
        r'^which\s+',
        r'^can\s+you',
        r'^could\s+you',
        r'^should\s+',
        r'^would\s+',
        r'^tell\s+me',
        r'^explain',
        r'^describe',
    ]

    for pattern in question_patterns:
        if re.match(pattern, message_lower):
            return True

    return False


def _extract_context_for_question(
    sections: List[dict],
    website_text: str,
    conversation_history: Optional[List[dict]] = None
) -> dict:
    """
    Extract full context for question answering:
    - Full document content (all sections)
    - Website summary (first 200-500 chars)
    - Conversation history (last 2-3 messages)
    """
    # Extract full document content
    document_content_parts = []
    for section in sections:
        section_id = section.get("id", "")
        section_title = section.get("title", "")
        section_content = section.get("content", "")
        if section_content and section_content.strip():
            document_content_parts.append(f"Section {section_id} ({section_title}): {section_content}")

    full_document_content = "\n\n".join(document_content_parts) if document_content_parts else "No content generated yet."

    # Extract website summary (200-500 chars)
    website_summary = ""
    if website_text:
        # Take first 500 chars, but try to end at a sentence boundary
        if len(website_text) > 500:
            # Find last sentence boundary within 500 chars
            truncated = website_text[:500]
            last_period = truncated.rfind('.')
            last_newline = truncated.rfind('\n')
            cut_point = max(last_period, last_newline)
            if cut_point > 200:  # Ensure we have at least 200 chars
                website_summary = website_text[:cut_point + 1]
            else:
                website_summary = website_text[:500]
        else:
            website_summary = website_text

    # Extract conversation history (last 2-3 messages)
    conversation_context = ""
    if conversation_history:
        # Take last 3 messages (or all if less than 3)
        recent_messages = conversation_history[-3:] if len(conversation_history) > 3 else conversation_history
        conversation_parts = []
        for msg in recent_messages:
            role = msg.get("role", "user")
            text = msg.get("text", "")
            if text:
                conversation_parts.append(f"{role.capitalize()}: {text}")
        conversation_context = "\n".join(conversation_parts) if conversation_parts else ""

    return {
        "document_content": full_document_content,
        "website_summary": website_summary,
        "conversation_history": conversation_context
    }


def _save_chat_message(
    document: Document,
    role: str,
    text: str,
    suggested_content: Optional[dict] = None,
    requires_confirmation: bool = False,
    db: Session = None
):
    """
    Save a chat message to the document's chat_history.
    """
    # Initialize chat_history if None
    if document.chat_history is None:
        document.chat_history = []
        logger.debug(f"Initialized chat_history for document {document.id}")

    # Create message object
    message = {
        "role": role,
        "text": text,
        "timestamp": datetime.now(timezone.utc).isoformat()
    }

    # Add optional fields
    if suggested_content:
        message["suggestedContent"] = suggested_content
    if requires_confirmation:
        message["requiresConfirmation"] = True
        message["messageId"] = f"msg-{int(datetime.now(timezone.utc).timestamp() * 1000)}"

    # Append to chat history
    document.chat_history.append(message)

    # Save to database
    try:
        db.commit()
        db.refresh(document)
        logger.info(f"Saved chat message to document {document.id}: role={role}, text_length={len(text)}, total_messages={len(document.chat_history)}")
    except Exception as e:
        logger.error(f"Failed to save chat message: {str(e)}", exc_info=True)
        db.rollback()
        # Don't raise - chat saving is not critical, but log the error


def _answer_question_with_context(
    client: OpenAI,
    user_query: str,
    document_content: str,
    website_summary: str,
    conversation_history: str,
    company_name: str
) -> str:
    """
    Answer a user question using full context (document, website, conversation history).
    Returns a concise answer in formal business language (Fördermittel tone).
    """
    # Build context prompt
    context_parts = []

    if document_content and document_content.strip() != "No content generated yet.":
        context_parts.append(f"Generated Document Content:\n{document_content}")

    if website_summary:
        context_parts.append(f"Company Website Summary:\n{website_summary}")

    if conversation_history:
        context_parts.append(f"Previous Conversation:\n{conversation_history}")

    context_text = "\n\n".join(context_parts)

    user_query_text = user_query or ""
    prompt = f"""Sie sind ein Expertenberater, der Fragen zu einem Förderantrag-Dokument (Vorhabensbeschreibung) beantwortet.

KONTEXT:
{context_text}

Firmenname: {company_name}

BENUTZERFRAGE: <user_instruction>
{user_query_text}
</user_instruction>

AUFGABE:
Beantworten Sie die Frage präzise und sachlich im formellen Fördermittel-Stil (Geschäftssprache).

WICHTIGE REGELN:
- Beziehen Sie sich AUSSCHLIESSLICH auf den bereitgestellten Kontext
- Wenn die Antwort nicht im Kontext enthalten ist, sagen Sie dies klar
- Verwenden Sie formelle, professionelle Sprache (Deutsch)
- Seien Sie präzise und konkret
- Keine Spekulationen oder Informationen außerhalb des Kontexts
- Keine Meta-Kommentare oder Hinweise auf KI
- Antworten Sie in zusammenhängenden Absätzen (keine Aufzählungen, außer wenn angebracht)

Geben Sie NUR die Antwort zurück, ohne zusätzliche Erklärungen oder Formatierungen."""

    approx_tokens = len(prompt) // 4
    logger.info("LLM Q&A prompt size (chars): %s", len(prompt))
    logger.info("LLM Q&A prompt tokens: %s", approx_tokens)
    try:
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {
                    "role": "system",
                    "content": "Sie sind ein professioneller Berater für Förderanträge. Sie beantworten Fragen präzise und sachlich im formellen Fördermittel-Stil."
                },
                {
                    "role": "user",
                    "content": prompt
                }
            ],
            temperature=0.7,
            max_tokens=1000,  # Limit response length for concise answers
            timeout=120.0  # 2 minute timeout for production safety
        )

        answer = response.choices[0].message.content.strip()
        logger.info(f"Generated answer for question: '{user_query[:50]}...' (answer length: {len(answer)})")
        return answer

    except Exception as e:
        logger.error(f"Error generating answer: {str(e)}")
        raise


@router.post(
    "/documents/{document_id}/chat",
    response_model=ChatResponse
)
def chat_with_document(
    document_id: int,
    chat_request: ChatRequest,
    db: Session = Depends(get_db),  # noqa: B008
    current_user: User = Depends(get_current_user)  # noqa: B008
):
    """
    ROLE: CHAT-BASED SECTION EDITING

    Chat endpoint for section-scoped editing of Vorhabensbeschreibung documents.
    Only calls LLM when user explicitly specifies section(s) and change instruction(s).
    Otherwise asks clarification questions.

    This endpoint:
    - Modifies existing section content based on user instructions
    - Uses _generate_section_content() for editing logic
    - Assumes sections already have content that needs modification

    This endpoint must NOT:
    - Call _generate_batch_content() (that's for initial generation only)
    - Be used for creating initial content (use /generate-content instead)
    """

    # Load document
    document = _safe_get_document_by_id(document_id, db)
    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found"
        )

    # Verify document type
    if document.type != "vorhabensbeschreibung":
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Chat editing only supported for vorhabensbeschreibung documents"
        )

    # Load associated company and verify ownership
    company = db.query(Company).filter(
        Company.id == document.company_id,
        Company.user_email == current_user.email
    ).first()
    if not company:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found"
        )

    # Load document sections
    content_json = document.content_json
    if not content_json or "sections" not in content_json:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Document has no sections"
        )

    sections = content_json["sections"]
    if not sections:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Document has no sections"
        )

    # Get valid section IDs
    valid_section_ids = [section.get("id", "") for section in sections if section.get("id")]

    # Get context (last edited sections) from request if available
    # Reserved for future use - not currently used in this function
    _last_edited_sections = chat_request.last_edited_sections
    conversation_history = chat_request.conversation_history or []

    # Check if message is a question
    is_question = _is_question(chat_request.message)

    if is_question:
        # Handle question-answering with full context
        logger.info(f"Detected question: '{chat_request.message[:50]}...'")

        # Extract context
        context = _extract_context_for_question(
            sections=sections,
            website_text=company.website_text or "",
            conversation_history=conversation_history
        )

        # Get OpenAI API key
        api_key = os.getenv("OPENAI_API_KEY")
        if not api_key:
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail="OpenAI API key not configured. Set OPENAI_API_KEY environment variable."
            )

        # Initialize OpenAI client
        try:
            client = OpenAI(api_key=api_key)
        except Exception as e:
            logger.error(f"Failed to initialize OpenAI client: {str(e)}")
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Failed to initialize OpenAI client: {str(e)}"
            ) from e

        # Answer the question with context
        try:
            answer = _answer_question_with_context(
                client=client,
                user_query=chat_request.message,
                document_content=context["document_content"],
                website_summary=context["website_summary"],
                conversation_history=context["conversation_history"],
                company_name=company.name or "Unknown Company"
            )

            logger.info(f"Question answered successfully (answer length: {len(answer)})")

            # Save user message and assistant response to chat history
            _save_chat_message(document, "user", chat_request.message, db=db)
            _save_chat_message(document, "assistant", answer, db=db)

            # Return answer without updating any sections
            # The frontend will display the answer in chat
            return ChatResponse(
                message=answer,
                updated_sections=None,
                is_question=True
            )

        except Exception as e:
            logger.error(f"Error answering question: {str(e)}")
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Failed to answer question: {str(e)}"
            ) from e

    # If not a question, proceed with section editing logic
    logger.info(f"Message is not a question - proceeding with section editing: '{chat_request.message[:50]}...'")

    # Save user message to chat history
    _save_chat_message(document, "user", chat_request.message, db=db)

    # Parse section changes: try enhanced parser first, fallback to original
    changes = _parse_section_changes_enhanced(chat_request.message, valid_section_ids, sections)

    # If enhanced parser found nothing, try original parser
    if not changes:
        changes = _parse_section_changes(chat_request.message, valid_section_ids, sections)

    # If still no changes found, create a default change with raw message
    # (This allows testing even with ambiguous requests)
    if not changes:
        logger.warning("No sections parsed from message, creating default change with first section")
        if valid_section_ids:
            # Use first section as default
            changes = [{
                "section_id": valid_section_ids[0],
                "instruction": chat_request.message
            }]
            logger.info(f"Created default change: section={valid_section_ids[0]}, instruction='{chat_request.message}'")
        else:
            return ChatResponse(
                message="Document has no sections to update.",
                updated_sections=None
            )

    # Validate changes (keep this for safety, but log warnings and continue)
    is_valid, error_msg = _validate_section_changes(changes, valid_section_ids)
    if not is_valid:
        logger.warning(f"Validation failed but proceeding anyway for testing: {error_msg}")
        # Continue anyway for testing - don't return error

    # Get OpenAI API key
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="OpenAI API key not configured. Set OPENAI_API_KEY environment variable."
        )

    # Initialize OpenAI client
    try:
        client = OpenAI(api_key=api_key)
    except Exception as e:
        logger.error(f"Failed to initialize OpenAI client: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to initialize OpenAI client: {str(e)}"
        ) from e

    # Prepare company data (use cleaned versions)
    company_name = company.name or "Unknown Company"
    website_clean_text = company.website_clean_text or None
    transcript_clean = company.transcript_clean or None
    company_profile = company.company_profile  # PRIMARY factual source (structured JSON)

    # Get style profile (system-level, from AlteVorhabensbeschreibung)
    style_profile = None
    from app.models import AlteVorhabensbeschreibungStyleProfile
    style_profile_record = db.query(AlteVorhabensbeschreibungStyleProfile).first()
    if style_profile_record:
        style_profile = style_profile_record.style_summary_json
        logger.info(f"Using style profile for chat editing (hash: {style_profile_record.combined_hash[:10]}...)")
    else:
        logger.warning("No style profile found for chat editing - using default style guidelines")

    # Process each section change
    updated_section_ids = []
    suggested_content_map = {}  # Map of section_id -> suggested_content for preview
    # Create a map for quick lookup, but we'll update the original sections list
    section_map = {section.get("id"): idx for idx, section in enumerate(sections)}

    for change in changes:
        section_id = change["section_id"]
        instruction = change["instruction"]

        if section_id not in section_map:
            logger.warning(f"Section {section_id} not found in document")
            continue

        section_idx = section_map[section_id]
        section = sections[section_idx]
        section_title = section.get("title", "")
        current_content = section.get("content", "")

        try:
            # Generate updated content
            # NOTE: This calls _generate_section_content (SECTION EDITOR role)
            # This is correct - we are editing existing content, not generating initial content
            logger.info(f"Calling LLM for section {section_id} with instruction: '{instruction}'")
            logger.info(f"Current content length: {len(current_content)} characters")
            logger.info(f"Current content preview: {current_content[:100] if current_content else '(empty)'}")

            new_content = _generate_section_content(
                client=client,
                section_id=section_id,
                section_title=section_title,
                current_content=current_content,
                instruction=instruction,
                company_name=company_name,
                company_profile=company_profile,  # PRIMARY factual source
                website_clean_text=website_clean_text,  # Contextual enrichment
                transcript_clean=transcript_clean,  # Contextual enrichment
                company_id=company.id,  # Guardrail A: Pass company_id for logging
                style_profile=style_profile  # Style guide
            )

            logger.info(f"LLM returned content length: {len(new_content)} characters")
            logger.info(f"LLM returned content preview: {new_content[:200]}")
            logger.info(f"Content changed: {new_content != current_content}")

            # Check if content actually changed (not just whitespace/formatting)
            content_changed = new_content.strip() != current_content.strip()
            if not content_changed:
                logger.warning(f"LLM returned identical content for section {section_id} - content was not actually modified!")
                logger.warning("This may indicate the LLM did not follow the rewrite/expand instruction properly")

            # Check if content is significantly longer (for expand/add instructions)
            length_increase = len(new_content) - len(current_content)
            length_increase_percent = (length_increase / len(current_content) * 100) if current_content else 0
            logger.info(f"Content length change: {length_increase} characters ({length_increase_percent:.1f}% increase)")

            # Store suggested content (DO NOT update section yet - wait for confirmation)
            # Build map of section_id -> suggested_content for preview
            suggested_content_map[section_id] = new_content
            updated_section_ids.append(section_id)
            logger.info(f"Successfully generated suggested content for section {section_id} for document {document_id} (preview mode)")

        except Exception as e:
            logger.error(f"Failed to generate content for section {section_id}: {str(e)}")
            # Continue with other sections even if one fails
            continue

    # Return preview instead of saving (user must confirm first)
    if updated_section_ids:
        # Generate preview response message
        try:
            if len(updated_section_ids) == 1:
                response_message = f"Ich habe eine Änderung für Abschnitt {updated_section_ids[0]} vorbereitet. Bitte überprüfen Sie die Vorschau und bestätigen Sie die Änderung."
            else:
                sections_str = ", ".join(updated_section_ids)
                response_message = f"Ich habe Änderungen für die Abschnitte {sections_str} vorbereitet. Bitte überprüfen Sie die Vorschau und bestätigen Sie die Änderungen."

            logger.info(f"Returning ChatResponse with preview for {len(updated_section_ids)} sections: {updated_section_ids}")

            # Save assistant response with preview to chat history
            _save_chat_message(
                document,
                "assistant",
                response_message,
                suggested_content=suggested_content_map,
                requires_confirmation=True,
                db=db
            )

            response = ChatResponse(
                message=response_message,
                suggested_content=suggested_content_map,
                requires_confirmation=True,
                updated_sections=None,  # Not updated yet - waiting for confirmation
                is_question=False  # Explicitly mark as section edit, not question
            )
            logger.info("ChatResponse with preview created successfully, returning...")
            return response
        except Exception as e:
            logger.error(f"Error creating ChatResponse: {str(e)}", exc_info=True)
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Failed to create response: {str(e)}"
            ) from e
    else:
        # No sections were updated (all failed)
        error_message = "Entschuldigung, es konnte kein Abschnitt aktualisiert werden. Bitte versuchen Sie es erneut mit spezifischeren Anweisungen."
        _save_chat_message(document, "assistant", error_message, db=db)
        return ChatResponse(
            message=error_message,
            updated_sections=None,
            is_question=False  # Explicitly mark as section edit attempt, not question
        )


@router.post(
    "/documents/{document_id}/chat/confirm",
    response_model=ChatResponse
)
def confirm_chat_edit(
    document_id: int,
    confirmation: ChatConfirmationRequest,
    db: Session = Depends(get_db),  # noqa: B008
    current_user: User = Depends(get_current_user)  # noqa: B008
):
    """
    Apply confirmed edit to a section.
    This endpoint is called when user approves a suggested edit from the preview.
    """

    # Load document
    document = _safe_get_document_by_id(document_id, db)
    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found"
        )

    # Verify company belongs to current user
    company = db.query(Company).filter(
        Company.id == document.company_id,
        Company.user_email == current_user.email
    ).first()
    if not company:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found"
        )

    # Verify document type
    if document.type != "vorhabensbeschreibung":
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Chat confirmation only supported for vorhabensbeschreibung documents"
        )

    # Load document sections
    content_json = document.content_json
    if not content_json or "sections" not in content_json:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Document has no sections"
        )

    sections = content_json["sections"]
    if not sections:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Document has no sections"
        )

    # Find the section to update
    section_found = False
    section_to_update = None
    logger.info(f"Looking for section {confirmation.section_id} in document {document_id}")
    logger.info(f"Available section IDs: {[s.get('id') for s in sections]}")

    for section in sections:
        section_id = section.get("id", "")
        if section_id == confirmation.section_id:
            section_to_update = section
            section_found = True
            break

    if not section_found or not section_to_update:
        logger.error(f"Section {confirmation.section_id} not found in document {document_id}. Available sections: {[s.get('id') for s in sections]}")
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=f"Section {confirmation.section_id} not found in document. Available sections: {', '.join([s.get('id', '') for s in sections])}"
        )

    # Only update content; never change section title.
    # If the LLM included the section title as the first line of suggested content, strip it so we don't duplicate the title.
    content_to_save = confirmation.confirmed_content
    section_title = (section_to_update.get("title") or "").strip()
    conf_section_id = confirmation.section_id or ""

    if content_to_save and section_title:
        lines = content_to_save.split("\n")
        first_line = (lines[0] or "").strip()
        # Strip first line only if it looks like the section heading (exact title or "1.3 Title" / "1.3. Title")
        strip_first = False
        if first_line == section_title:
            strip_first = True
        elif first_line == f"{conf_section_id} {section_title}" or first_line == f"{conf_section_id}. {section_title}":
            strip_first = True
        elif first_line.endswith(section_title):
            rest = first_line[: -len(section_title)].strip().rstrip(".")
            if rest == conf_section_id or not rest:
                strip_first = True
        # e.g. section_title "1.3. Branche und Leistungsangebot", LLM first line "1.3 Branche und Leistungsangebot"
        elif conf_section_id and first_line.startswith(conf_section_id):
            after_id = first_line[len(conf_section_id):].strip().lstrip(".").strip()
            if after_id == section_title or section_title.endswith(after_id) or after_id == section_title.replace(f"{conf_section_id}. ", "", 1):
                strip_first = True
        if strip_first and len(lines) > 1:
            content_to_save = "\n".join(lines[1:]).strip()
        elif strip_first:
            content_to_save = ""

    section_to_update["content"] = content_to_save
    logger.info(f"Updating section {confirmation.section_id} with confirmed content (title unchanged, content length: {len(content_to_save)})")

    # Rebuild sections array preserving order
    # IMPORTANT: Verify the updated content is in the sections list before rebuilding
    updated_sections = []
    for section in sections:
        section_id = section.get("id", "")
        section_content = section.get("content", "")

        # Log if this is the section we just updated
        if section_id == confirmation.section_id:
            logger.info(f"Rebuilding section {section_id} with content length: {len(section_content)} (expected: {len(content_to_save)})")
            if section_content != content_to_save:
                logger.error(f"ERROR: Section {section_id} content mismatch during rebuild! Setting correct content.")
                section_content = content_to_save  # Force correct content

        section_data = {
            "id": section_id,
            "title": section.get("title", ""),
            "content": section_content
        }
        if section.get("type") is not None:
            section_data["type"] = section.get("type")
        updated_sections.append(section_data)

    # Verify the updated section is in the rebuilt array
    rebuilt_section = next((s for s in updated_sections if s.get("id") == confirmation.section_id), None)
    if rebuilt_section:
        logger.info(f"Rebuilt section {confirmation.section_id} content length: {len(rebuilt_section.get('content', ''))}")
        if rebuilt_section.get("content") != content_to_save:
            logger.error("ERROR: Rebuilt section content doesn't match! Forcing correct content.")
            rebuilt_section["content"] = content_to_save

    # Update document in database
    # IMPORTANT: Create a new dict to ensure SQLAlchemy detects the change
    document.content_json = {"sections": updated_sections}

    # Mark the JSON column as modified so SQLAlchemy knows to update it
    # This is REQUIRED when modifying nested JSON structures - SQLAlchemy doesn't detect nested changes automatically
    flag_modified(document, "content_json")

    try:
        db.commit()
        db.refresh(document)

        # Verify the content was actually saved
        saved_section = None
        for s in document.content_json.get("sections", []):
            if s.get("id") == confirmation.section_id:
                saved_section = s
                break

        if saved_section:
            saved_content = saved_section.get("content", "")
            logger.info(f"Successfully saved confirmed edit for section {confirmation.section_id} in document {document_id}")
            logger.info(f"Verified saved content length: {len(saved_content)} (expected: {len(content_to_save)})")
            if saved_content != content_to_save:
                logger.error("ERROR: Saved content does not match confirmed content!")
                logger.error(f"Expected preview: {content_to_save[:200]}...")
                logger.error(f"Got preview: {saved_content[:200]}...")
                # This is a critical error - raise an exception
                raise HTTPException(
                    status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                    detail="Failed to save content: content mismatch after save"
                )
            else:
                logger.info("✓ Content verified successfully - saved content matches confirmed content")
        else:
            logger.error(f"ERROR: Section {confirmation.section_id} not found in saved document!")

    except Exception as e:
        db.rollback()
        logger.error(f"Failed to save confirmed edit for document {document_id}: {str(e)}", exc_info=True)
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to save confirmed edit: {str(e)}"
        ) from e

    # Return success response
    return ChatResponse(
        message=f"Änderung für Abschnitt {confirmation.section_id} wurde bestätigt und gespeichert.",
        updated_sections=[confirmation.section_id],
        is_question=False,
        requires_confirmation=False
    )


@router.get("/documents/{document_id}/export")
def export_document(
    document_id: int,
    format: str = "pdf",  # "pdf" or "docx"
    db: Session = Depends(get_db),  # noqa: B008
    current_user: User = Depends(get_current_user)  # noqa: B008
):
    """
    Export document as PDF or DOCX file.
    """

    # Load document
    document = _safe_get_document_by_id(document_id, db)
    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found"
        )

    # Load associated company for filename and verify ownership
    company = db.query(Company).filter(
        Company.id == document.company_id,
        Company.user_email == current_user.email
    ).first()
    if not company:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found"
        )
    company_name = company.name
    # Sanitize filename
    safe_company_name = re.sub(r'[^\w\s-]', '', company_name).strip().replace(' ', '_')

    # Get document content
    content_json = document.content_json
    if not content_json or "sections" not in content_json:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Document has no content to export"
        )

    sections = content_json["sections"]

    if format.lower() == "pdf":
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
            from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT
            from reportlab.lib import colors
            from reportlab.lib.units import mm

            # Create PDF in memory
            buffer = io.BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=A4)
            story = []
            styles = getSampleStyleSheet()

            # Custom styles
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=14,
                textColor=(0.2, 0.2, 0.2),
                spaceAfter=12,
                fontName='Helvetica-Bold'
            )

            content_style = ParagraphStyle(
                'CustomContent',
                parent=styles['Normal'],
                fontSize=11,
                textColor=(0.1, 0.1, 0.1),
                spaceAfter=12,
                leftIndent=0,
                alignment=TA_LEFT
            )

            # Add sections to PDF
            for section in sections:
                title = section.get("title", "")
                content = section.get("content", "")
                section_type = section.get("type", "text")

                if title:
                    story.append(Paragraph(title, title_style))
                    story.append(Spacer(1, 6))

                # Handle milestone tables
                if section_type == "milestone_table":
                    try:
                        # Parse milestone JSON
                        if isinstance(content, str) and content.strip():
                            milestone_data = json.loads(content)
                        elif isinstance(content, dict):
                            milestone_data = content
                        else:
                            milestone_data = {"milestones": [], "total_expenditure": None}

                        milestones = milestone_data.get("milestones", [])
                        total_expenditure = milestone_data.get("total_expenditure", None)

                        if milestones:
                            # Create table data
                            table_data = []
                            
                            # Header row
                            header_row = [
                                Paragraph("<b>Meilenstein</b>", content_style),
                                Paragraph("<b>erwartetes Ziel</b>", content_style),
                                Paragraph("<b>erwarteter Zeitpunkt der Zielerreichung<br/>(TT.MM.JJJJ)</b>", content_style),
                                Paragraph("<b>erwartete Ausgaben zum Zeitpunkt<br/>der Zielerreichung (EUR)</b>", content_style)
                            ]
                            table_data.append(header_row)

                            # Helper function to format numbers in German format (1 500,08)
                            def format_german_number(value: float) -> str:
                                """Format number as German format: 1 500,08"""
                                formatted = f"{value:,.2f}"  # US format: 1,500.08
                                # Replace thousands comma with space, then decimal period with comma
                                return formatted.replace(",", "X").replace(".", ",").replace("X", " ")

                            # Data rows
                            for milestone in milestones:
                                expenditure = milestone.get("expected_expenditure", 0)
                                row = [
                                    Paragraph(str(milestone.get("milestone_number", "")), content_style),
                                    Paragraph(str(milestone.get("expected_target", "")), content_style),
                                    Paragraph(str(milestone.get("target_date", "")), content_style),
                                    Paragraph(format_german_number(expenditure), content_style)
                                ]
                                table_data.append(row)

                            # Total row
                            total_value = total_expenditure if total_expenditure is not None else sum(
                                m.get("expected_expenditure", 0) for m in milestones
                            )
                            total_row = [
                                "",
                                "",
                                Paragraph("<b>erwartete Gesamtausgaben</b>", content_style),
                                Paragraph(f"<b>{format_german_number(total_value)}</b>", content_style)
                            ]
                            table_data.append(total_row)

                            # Create table with proper column widths
                            col_widths = [20*mm, 60*mm, 50*mm, 50*mm]
                            table = Table(table_data, colWidths=col_widths, repeatRows=1)

                            # Style the table
                            table.setStyle(TableStyle([
                                # Header styling
                                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                                ('ALIGN', (3, 1), (3, -1), 'RIGHT'),  # Right-align expenditure column
                                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                                ('FONTSIZE', (0, 0), (-1, 0), 10),
                                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                                ('TOPPADDING', (0, 0), (-1, 0), 12),
                                
                                # Grid lines
                                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                                ('LINEBELOW', (0, 0), (-1, 0), 2, colors.black),
                                
                                # Data row styling
                                ('FONTNAME', (0, 1), (-1, -2), 'Helvetica'),
                                ('FONTSIZE', (0, 1), (-1, -2), 10),
                                ('ROWBACKGROUNDS', (0, 1), (-1, -2), [colors.white, colors.lightgrey]),
                                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                                ('LEFTPADDING', (0, 0), (-1, -1), 6),
                                ('RIGHTPADDING', (0, 0), (-1, -1), 6),
                                ('TOPPADDING', (0, 1), (-1, -2), 8),
                                ('BOTTOMPADDING', (0, 1), (-1, -2), 8),
                                
                                # Total row styling
                                ('FONTNAME', (2, -1), (3, -1), 'Helvetica-Bold'),
                                ('FONTSIZE', (2, -1), (3, -1), 10),
                                ('LINEABOVE', (0, -1), (-1, -1), 2, colors.black),
                                ('TOPPADDING', (0, -1), (-1, -1), 8),
                                ('BOTTOMPADDING', (0, -1), (-1, -1), 8),
                            ]))

                            story.append(table)
                            story.append(Spacer(1, 12))
                        else:
                            # Empty milestone table
                            story.append(Paragraph("Keine Meilensteine definiert.", content_style))
                            story.append(Spacer(1, 12))
                    except (json.JSONDecodeError, KeyError, TypeError) as e:
                        logger.warning(f"Failed to parse milestone table for section {section.get('id', 'unknown')}: {str(e)}")
                        # Fallback to text representation
                        content_str = str(content) if content else ""
                        content_escaped = content_str.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                        story.append(Paragraph(content_escaped, content_style))
                        story.append(Spacer(1, 12))
                else:
                    # Regular text section
                    # Ensure content is a string (handle dict/other types)
                    if not isinstance(content, str):
                        if isinstance(content, dict):
                            # If content is a dict, convert to string representation
                            content = str(content)
                        elif content is None:
                            content = ""
                        else:
                            content = str(content)

                    if content:
                        # Escape HTML and convert newlines
                        content_escaped = content.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                        content_escaped = content_escaped.replace("\n", "<br/>")
                        story.append(Paragraph(content_escaped, content_style))
                        story.append(Spacer(1, 12))

            # Build PDF
            doc.build(story)
            buffer.seek(0)

            filename = f"{safe_company_name}_Vorhabensbeschreibung.pdf"
            return Response(
                content=buffer.getvalue(),
                media_type="application/pdf",
                headers={
                    "Content-Disposition": f"attachment; filename={filename}"
                }
            )
        except ImportError:
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail="PDF export requires reportlab library. Install with: pip install reportlab"
            ) from None
        except Exception as e:
            logger.error(f"PDF export error for document {document_id}: {str(e)}")
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Failed to generate PDF: {str(e)}"
            ) from e

    elif format.lower() == "docx" or format.lower() == "doc":
        try:
            from docx import Document as DocxDocument
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            # Create DOCX document
            docx = DocxDocument()

            # Add sections to DOCX
            for section in sections:
                title = section.get("title", "")
                content = section.get("content", "")
                section_type = section.get("type", "text")

                if title:
                    title_para = docx.add_paragraph(title)
                    title_para.style = 'Heading 1'
                    title_run = title_para.runs[0] if title_para.runs else title_para.add_run(title)
                    title_run.font.size = Pt(14)
                    title_run.bold = True

                # Handle milestone tables
                if section_type == "milestone_table":
                    try:
                        # Parse milestone JSON
                        if isinstance(content, str) and content.strip():
                            milestone_data = json.loads(content)
                        elif isinstance(content, dict):
                            milestone_data = content
                        else:
                            milestone_data = {"milestones": [], "total_expenditure": None}

                        milestones = milestone_data.get("milestones", [])
                        total_expenditure = milestone_data.get("total_expenditure", None)

                        if milestones:
                            # Create table
                            table = docx.add_table(rows=1, cols=4)
                            table.style = 'Light Grid Accent 1'

                            # Header row
                            header_cells = table.rows[0].cells
                            header_cells[0].text = "Meilenstein"
                            header_cells[1].text = "erwartetes Ziel"
                            header_cells[2].text = "erwarteter Zeitpunkt der Zielerreichung (TT.MM.JJJJ)"
                            header_cells[3].text = "erwartete Ausgaben zum Zeitpunkt der Zielerreichung (EUR)"

                            # Make header bold
                            for cell in header_cells:
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.bold = True
                                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

                            # Helper function to format numbers in German format (1 500,08)
                            def format_german_number(value: float) -> str:
                                """Format number as German format: 1 500,08"""
                                formatted = f"{value:,.2f}"  # US format: 1,500.08
                                # Replace thousands comma with space, then decimal period with comma
                                return formatted.replace(",", "X").replace(".", ",").replace("X", " ")

                            # Add data rows
                            for milestone in milestones:
                                row_cells = table.add_row().cells
                                row_cells[0].text = str(milestone.get("milestone_number", ""))
                                row_cells[1].text = str(milestone.get("expected_target", ""))
                                row_cells[2].text = str(milestone.get("target_date", ""))
                                
                                # Format expenditure with German number format
                                expenditure = milestone.get("expected_expenditure", 0)
                                row_cells[3].text = format_german_number(expenditure)
                                
                                # Right-align expenditure column
                                row_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

                            # Add total row
                            total_value = total_expenditure if total_expenditure is not None else sum(
                                m.get("expected_expenditure", 0) for m in milestones
                            )
                            total_row = table.add_row()
                            total_cells = total_row.cells
                            total_cells[0].text = ""
                            total_cells[1].text = ""
                            total_cells[2].text = "erwartete Gesamtausgaben"
                            total_cells[3].text = format_german_number(total_value)

                            # Make total row bold
                            for cell in [total_cells[2], total_cells[3]]:
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.bold = True
                                    if cell == total_cells[3]:
                                        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

                            # Add spacing after table
                            docx.add_paragraph()
                        else:
                            # Empty milestone table
                            empty_para = docx.add_paragraph("Keine Meilensteine definiert.")
                            empty_para.style = 'Normal'
                            for run in empty_para.runs:
                                run.font.size = Pt(11)
                            docx.add_paragraph()
                    except (json.JSONDecodeError, KeyError, TypeError) as e:
                        logger.warning(f"Failed to parse milestone table for section {section.get('id', 'unknown')}: {str(e)}")
                        # Fallback to text representation
                        content_str = str(content) if content else ""
                        content_para = docx.add_paragraph(content_str)
                        content_para.style = 'Normal'
                        for run in content_para.runs:
                            run.font.size = Pt(11)
                        docx.add_paragraph()
                else:
                    # Regular text section
                    if content:
                        content_para = docx.add_paragraph(content)
                        content_para.style = 'Normal'
                        for run in content_para.runs:
                            run.font.size = Pt(11)
                        # Add spacing after content
                        docx.add_paragraph()

            # Save to buffer
            buffer = io.BytesIO()
            docx.save(buffer)
            buffer.seek(0)

            filename = f"{safe_company_name}_Vorhabensbeschreibung.docx"
            return Response(
                content=buffer.getvalue(),
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={
                    "Content-Disposition": f"attachment; filename={filename}"
                }
            )
        except ImportError:
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail="DOCX export requires python-docx library. Install with: pip install python-docx"
            ) from None
        except Exception as e:
            logger.error(f"DOCX export error for document {document_id}: {str(e)}")
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Failed to generate DOCX: {str(e)}"
            ) from e

    else:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Unsupported export format: {format}. Supported formats: pdf, docx"
        )

